# =================================================================
# برنامج إدارة مداخيل ومصاريف جمعية المسجد - الإصدار 8.5 (إصلاح خطأ الرسوم البيانية)
# =================================================================
# المكتبات المطلوبة:
# pip install customtkinter
# pip install pandas openpyxl numpy
# pip install Pillow
# pip install matplotlib arabic-reshaper python-bidi
# pip install python-docx
# =================================================================
# التحسينات الجديدة في هذا الإصدار (v8.5):
# - (إصلاح خطأ) معالجة خطأ ValueError عند إنشاء الرسوم البيانية في قسم التقارير بسبب عدم توافق الألوان.
# - (ميزة جديدة) إعادة هيكلة قسم التقارير بالكامل.
# - (ميزة جديدة) إضافة "تقارير التبرعات" (ملخص، مالية، عينية).
# - (ميزة جديدة) إضافة "تقارير الأداء" (شهري، سنوي).
# - (ميزة جديدة) إضافة قسم "للتنبؤ المالي" المستقبلي.
# - إضافة فئة "تبرعات عينية" للمداخيل.
# =================================================================

import customtkinter as ctk
import sqlite3
from tkinter import ttk, messagebox, filedialog
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import hashlib
import os
import sys
import traceback
import shutil
import subprocess
from openpyxl.styles import PatternFill, Font as OpenpyxlFont, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageTk, ImageFont
import matplotlib
matplotlib.use("TkAgg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import matplotlib.font_manager as fm
import arabic_reshaper
from bidi.algorithm import get_display
# --- إضافة المكتبات المطلوبة لملفات الوورد ---
try:
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    messagebox.showerror(
        "مكتبة ناقصة",
        "مكتبة 'python-docx' غير مثبتة.\n"
        "لإنشاء الوصولات، يرجى تثبيتها عبر الأمر التالي:\n\n"
        "pip install python-docx"
    )
    sys.exit(1)


# =================================================================
# إعدادات المظهر والخطوط
# =================================================================
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")
APP_VERSION = "8.5"

# --- مدير الخطوط المركزي ---
class FontManager:
    """
    فئة مركزية لإدارة تحميل والوصول إلى الخطوط المخصصة.
    تستخدم "Cairo" للواجهة و "Amiri" للرسوم البيانية لضمان أفضل توافق.
    """
    FONT_DIR = "fonts"
    UI_FONT_FAMILY = "Cairo"
    CHART_FONT_FAMILY = "Amiri"
    BODY_FONT_FAMILY = "Tajawal"

    FONT_FILES = {
        "Cairo-Regular": "Cairo-Regular.ttf",
        "Cairo-Bold": "Cairo-Bold.ttf",
        "Cairo-ExtraBold": "Cairo-ExtraBold.ttf",
        "Cairo-Black": "Cairo-Black.ttf",
        "Cairo-ExtraLight": "Cairo-ExtraLight.ttf",
        "Cairo-Light": "Cairo-Light.ttf",
        "Cairo-Medium": "Cairo-Medium.ttf",
        "Cairo-SemiBold": "Cairo-SemiBold.ttf",
        "Amiri-Regular": "Amiri-Regular.ttf",
        "Amiri-Bold": "Amiri-Bold.ttf",
        "Tajawal-Black": "Tajawal-Black.ttf",
        "Tajawal-Bold": "Tajawal-Bold.ttf",
        "Tajawal-ExtraBold": "Tajawal-ExtraBold.ttf",
        "Tajawal-ExtraLight": "Tajawal-ExtraLight.ttf",
        "Tajawal-Light": "Tajawal-Light.ttf",
        "Tajawal-Medium": "Tajawal-Medium.ttf",
        "Tajawal-Regular": "Tajawal-Regular.ttf",

    }

    TITLE_FONT, H1_FONT, H2_FONT, BUTTON_FONT, APP_FONT, INPUT_FONT, TABLE_HEADER_FONT, TABLE_BODY_FONT, SMALL_FONT, BODY_FONT_FAMILY, = (None,) * 10
    _font_paths_cache = {}
    _matplotlib_fonts_cache = {}

    @staticmethod
    def get_font_path(font_family, weight="Regular"):
        font_key = f"{font_family}-{weight}"
        if font_key in FontManager._font_paths_cache:
            return FontManager._font_paths_cache[font_key]

        try:
            base_path = sys._MEIPASS
        except AttributeError:
            base_path = os.path.abspath(".")

        font_file = FontManager.FONT_FILES.get(font_key)
        if not font_file:
            font_file = FontManager.FONT_FILES.get(f"{font_family}-Regular", FontManager.FONT_FILES["Cairo-Regular"])

        path = os.path.join(base_path, FontManager.FONT_DIR, font_file)
        FontManager._font_paths_cache[font_key] = path
        return path

    @staticmethod
    def initialize_fonts():
        """Initializes CTkFont objects for the UI using the Cairo font family."""
        FontManager.TITLE_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=26, weight="bold")
        FontManager.H1_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=20, weight="bold")
        FontManager.H2_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=18, weight="bold")
        FontManager.BUTTON_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=16, weight="bold")
        FontManager.APP_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=15)
        FontManager.INPUT_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=15)
        FontManager.SMALL_FONT = ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=12)
        FontManager.TABLE_HEADER_FONT = (FontManager.UI_FONT_FAMILY, 16, "bold")
        FontManager.TABLE_BODY_FONT = (FontManager.UI_FONT_FAMILY, 14)

    @staticmethod
    def get_matplotlib_font_prop(weight="Regular", size=12):
        """Returns FontProperties for Matplotlib, defaulting to the defined CHART_FONT_FAMILY."""
        cache_key = (FontManager.CHART_FONT_FAMILY, weight, size)
        if cache_key in FontManager._matplotlib_fonts_cache:
            return FontManager._matplotlib_fonts_cache[cache_key]

        font_path = FontManager.get_font_path(FontManager.CHART_FONT_FAMILY, weight)
        if not os.path.exists(font_path):
            font_path = FontManager.get_font_path(FontManager.UI_FONT_FAMILY, "Regular")

        prop = fm.FontProperties(fname=font_path, size=size)
        FontManager._matplotlib_fonts_cache[cache_key] = prop
        return prop

    @staticmethod
    def check_and_register_fonts():
        """Checks for all required font files and registers them."""
        missing_fonts = []
        required_font_keys = [
            "Cairo-Regular", "Cairo-Bold",
            "Amiri-Regular", "Amiri-Bold"
        ]

        for font_key in required_font_keys:
            path = FontManager.get_font_path(font_key.split('-')[0], font_key.split('-')[1])
            if not os.path.exists(path):
                missing_fonts.append(FontManager.FONT_FILES.get(font_key, f"{font_key}.ttf"))

        if missing_fonts:
            return (f"ملفات الخطوط التالية مفقودة من مجلد 'fonts':\n\n"
                    f"{', '.join(missing_fonts)}\n\n"
                    "الرجاء التأكد من تحميل ووضع جميع ملفات خطوط 'Cairo' و 'Amiri' المطلوبة في مجلد 'fonts'.")

        for font_key, filename in FontManager.FONT_FILES.items():
            font_family, weight = font_key.split('-')
            path = FontManager.get_font_path(font_family, weight)
            if os.path.exists(path):
                try:
                    fm.fontManager.addfont(path)
                except Exception as e:
                    print(f"فشل تسجيل الخط {path} مع matplotlib: {e}")

        try:
             fm._load_fontmanager(try_read_cache=False)
        except Exception as e:
            print(f"فشل إعادة بناء ذاكرة خطوط matplotlib: {e}")

        plt.rcParams['font.family'] = FontManager.CHART_FONT_FAMILY
        plt.rcParams['axes.unicode_minus'] = False
        return None

# =================================================================
# دالة تنسيق النص العربي (للرسوم البيانية)
# =================================================================
def format_arabic(text):
    return get_display(arabic_reshaper.reshape(str(text)))

# =================================================================
# دالة التفقيط (تحويل الأرقام إلى نصوص عربية)
# =================================================================
def tafqeet(number):
    if number is None:
        return "غير محدد"
    try:
        number = int(number)
    except (ValueError, TypeError):
        return "قيمة غير صالحة"
    if number == 0:
        return "صفر"

    to_19 = ["", "واحد", "اثنان", "ثلاثة", "أربعة", "خمسة", "ستة", "سبعة", "ثمانية", "تسعة",
             "عشرة", "أحد عشر", "اثنا عشر", "ثلاثة عشر", "أربعة عشر", "خمسة عشر",
             "ستة عشر", "سبعة عشر", "ثمانية عشر", "تسعة عشر"]

    tens = ["", "", "عشرون", "ثلاثون", "أربعون", "خمسون", "ستون", "سبعون", "ثمانون", "تسعون"]
    hundreds = ["", "مئة", "مئتان", "ثلاثمئة", "أربعمئة", "خمسمئة", "ستمئة", "سبعمئة", "ثمانمئة", "تسعمئة"]
    thousands = ["ألف", "ألفان", "آلاف"]
    millions = ["مليون", "مليونان", "ملايين"]

    def convert_less_than_thousand(n):
        if n == 0:
            return ""
        if n < 20:
            return to_19[n]
        elif n < 100:
            ten, unit = divmod(n, 10)
            if unit == 0:
                return tens[ten]
            return f"{to_19[unit]} و{tens[ten]}"
        else:
            hundred, rest = divmod(n, 100)
            base = hundreds[hundred]
            if rest == 0:
                return base
            return f"{base} و{convert_less_than_thousand(rest)}"

    result_parts = []

    if number >= 1_000_000:
        millions_part = number // 1_000_000
        number %= 1_000_000
        if millions_part == 1:
            result_parts.append(millions[0])
        elif millions_part == 2:
            result_parts.append(millions[1])
        elif 3 <= millions_part <= 10:
            result_parts.append(f"{convert_less_than_thousand(millions_part)} {millions[2]}")
        else:
            result_parts.append(f"{convert_less_than_thousand(millions_part)} {millions[0]}")

    if number >= 1000:
        thousands_part = number // 1000
        number %= 1000
        if thousands_part == 1:
            result_parts.append(thousands[0])
        elif thousands_part == 2:
            result_parts.append(thousands[1])
        elif 3 <= thousands_part <= 10:
            result_parts.append(f"{convert_less_than_thousand(thousands_part)} {thousands[2]}")
        else:
            result_parts.append(f"{convert_less_than_thousand(thousands_part)} {thousands[0]}")

    if number > 0:
        result_parts.append(convert_less_than_thousand(number))

    return " و".join(result_parts)


# =================================================================
# مدير الأيقونات (IconManager)
# =================================================================
class IconManager:
    @staticmethod
    def create_icon(draw_function, size=(24, 24)):
        light_image = Image.new("RGBA", size, (0, 0, 0, 0))
        draw_light = ImageDraw.Draw(light_image)
        draw_function(draw_light, "#1C1C1C")

        dark_image = Image.new("RGBA", size, (0, 0, 0, 0))
        draw_dark = ImageDraw.Draw(dark_image)
        draw_function(draw_dark, "#DCE4EE")

        return ctk.CTkImage(light_image=light_image, dark_image=dark_image, size=size)

    @staticmethod
    def load_all_icons():
        icons = {}
        icons["dashboard"] = IconManager.create_icon(lambda d, c: (d.rectangle((3, 3, 10, 10), fill=c), d.rectangle((14, 3, 21, 10), fill=c), d.rectangle((3, 14, 10, 21), fill=c), d.rectangle((14, 14, 21, 21), fill=c)))
        icons["income"] = IconManager.create_icon(lambda d, c: (d.line((12, 4, 12, 20), fill=c, width=2), d.line((8, 16, 12, 20), fill=c, width=2), d.line((16, 16, 12, 20), fill=c, width=2), d.text((7, 2), "+", fill=c, font=ImageFont.load_default().font_variant(size=20))))
        icons["expense"] = IconManager.create_icon(lambda d, c: (d.line((12, 4, 12, 20), fill=c, width=2), d.line((8, 8, 12, 4), fill=c, width=2), d.line((16, 8, 12, 4), fill=c, width=2), d.text((8, 12), "-", fill=c, font=ImageFont.load_default().font_variant(size=20))))
        icons["reports"] = IconManager.create_icon(lambda d, c: (d.rectangle((4, 4, 20, 20), outline=c, width=2), d.line((8, 9, 16, 9), fill=c, width=1), d.line((8, 13, 16, 13), fill=c, width=1), d.line((8, 17, 12, 17), fill=c, width=1)))
        icons["data"] = IconManager.create_icon(lambda d, c: (d.ellipse((4, 4, 20, 12), outline=c, width=2), d.line((4, 8, 20, 8), fill=c, width=2), d.line((10, 12, 10, 20), fill=c, width=2), d.line((14, 12, 14, 20), fill=c, width=2)))
        icons["users"] = IconManager.create_icon(lambda d, c: (d.ellipse((8, 4, 16, 12), outline=c, width=2), d.arc((4, 12, 20, 24), 20, 160, fill=c, width=2)))
        icons["audit"] = IconManager.create_icon(lambda d, c: (d.polygon([(6, 4), (18, 4), (18, 20), (14, 20), (14, 8), (6, 8), (6, 20), (10,20), (10,14), (6,14)], outline=c, width=2)))
        icons["logout"] = IconManager.create_icon(lambda d, c: (d.rectangle((4, 4, 14, 20), outline=c, width=2), d.line((14, 12, 20, 12), fill=c, width=2), d.line((17, 9, 20, 12), fill=c, width=2), d.line((17, 15, 20, 12), fill=c, width=2)))
        icons["add"] = IconManager.create_icon(lambda d, c: (d.line((12, 5, 12, 19), fill=c, width=2), d.line((5, 12, 19, 12), fill=c, width=2)))
        icons["edit"] = IconManager.create_icon(lambda d, c: d.polygon([(5, 19), (5, 15), (16, 4), (20, 8), (9, 19), (5, 19)], outline=c, width=2))
        icons["delete"] = IconManager.create_icon(lambda d, c: (d.rectangle((6, 6, 18, 20), outline=c, width=2), d.line((4, 4, 20, 4), fill=c, width=2), d.line((10, 10, 10, 16), fill=c, width=2), d.line((14, 10, 14, 16), fill=c, width=2)))
        icons["export"] = IconManager.create_icon(lambda d, c: (d.rectangle((4,8,20,20),outline=c, width=2), d.line((12,2,12,14),fill=c, width=2), d.line((8,10,12,14),fill=c, width=2), d.line((16,10,12,14),fill=c, width=2)))
        icons["import"] = IconManager.create_icon(lambda d, c: (d.rectangle((4,14,20,20),outline=c, width=2), d.line((12,2,12,16),fill=c, width=2), d.line((8,8,12,2),fill=c, width=2), d.line((16,8,12,2),fill=c, width=2)))
        icons["search"] = IconManager.create_icon(lambda d, c: (d.ellipse((4, 4, 16, 16), outline=c, width=2), d.line((14, 14, 20, 20), fill=c, width=2)))
        icons["about"] = IconManager.create_icon(lambda d, c: (d.ellipse((4,4,20,20), outline=c, width=2), d.text((11, 4), "i", fill=c, font=ImageFont.load_default().font_variant(size=18))))
        icons["attachment"] = IconManager.create_icon(lambda d, c: (d.polygon([(6,2), (18,2), (18,22), (6,22)], outline=c, width=2), d.ellipse((7,4,11,8), fill=c)))
        icons["category"] = IconManager.create_icon(lambda d, c: (d.rectangle((3, 8, 21, 16), outline=c, width=2), d.line((7,8,7,4),fill=c,width=2), d.line((12,8,12,4),fill=c,width=2), d.line((17,8,17,4),fill=c,width=2)))
        icons["print"] = IconManager.create_icon(lambda d, c: (d.rectangle((4, 16, 20, 20), outline=c, width=2), d.rectangle((7, 8, 17, 16), outline=c, width=2), d.line((4,12,7,12), fill=c, width=2), d.rectangle((7,4,17,8),fill=c)))
        icons["activity"] = IconManager.create_icon(lambda d, c: (d.rectangle((4, 6, 20, 20), outline=c, width=2), d.line((8, 4, 8, 8), fill=c, width=2), d.line((16, 4, 16, 8), fill=c, width=2), d.line((4, 10, 20, 10), fill=c, width=2)))
        return icons

# =================================================================
# الفئة الخاصة بإدارة قاعدة البيانات (DatabaseManager)
# =================================================================
class DatabaseManager:
    def __init__(self, db_name="masjid_pro_database_v6.db"):
        self.conn = sqlite3.connect(db_name)
        self.conn.row_factory = sqlite3.Row
        self.cursor = self.conn.cursor()
        self.setup_tables()

    def hash_password(self, password):
        return hashlib.sha256(password.encode()).hexdigest()

    def setup_tables(self):
        self.cursor.execute("CREATE TABLE IF NOT EXISTS incomes (id INTEGER PRIMARY KEY, amount REAL NOT NULL, date TEXT NOT NULL, category TEXT NOT NULL, description TEXT, notes TEXT, payer TEXT, attachment_path TEXT)")
        self.cursor.execute("CREATE TABLE IF NOT EXISTS expenses (id INTEGER PRIMARY KEY, amount REAL NOT NULL, date TEXT NOT NULL, category TEXT NOT NULL, description TEXT, notes TEXT, attachment_path TEXT)")
        self.cursor.execute("CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY, username TEXT UNIQUE NOT NULL, password TEXT NOT NULL, role TEXT NOT NULL, must_change_password INTEGER DEFAULT 0)")
        self.cursor.execute("CREATE TABLE IF NOT EXISTS audit_log (id INTEGER PRIMARY KEY, timestamp TEXT NOT NULL, username TEXT NOT NULL, action TEXT NOT NULL, details TEXT)")
        self.cursor.execute("CREATE TABLE IF NOT EXISTS income_categories (id INTEGER PRIMARY KEY, name TEXT UNIQUE NOT NULL)")
        self.cursor.execute("CREATE TABLE IF NOT EXISTS expense_categories (id INTEGER PRIMARY KEY, name TEXT UNIQUE NOT NULL)")
        
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS members (
                id INTEGER PRIMARY KEY,
                full_name TEXT NOT NULL,
                join_date TEXT NOT NULL,
                phone TEXT,
                address TEXT,
                status TEXT DEFAULT 'نشط',
                notes TEXT
            )
        """)
        
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS activities (
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                date TEXT NOT NULL,
                location TEXT,
                description TEXT
            )
        """)
        
        self.cursor.execute("""
            CREATE TABLE IF NOT EXISTS activity_attendance (
                id INTEGER PRIMARY KEY,
                activity_id INTEGER NOT NULL,
                member_id INTEGER NOT NULL,
                FOREIGN KEY (activity_id) REFERENCES activities(id) ON DELETE CASCADE,
                FOREIGN KEY (member_id) REFERENCES members(id) ON DELETE CASCADE,
                UNIQUE(activity_id, member_id)
            )
        """)
        
        self.setup_settings_table()

        self.cursor.execute("SELECT COUNT(*) FROM users")
        if self.cursor.fetchone()[0] == 0:
            self.add_user("admin", "admin", "مدير", must_change_password=1)
        self.init_default_categories()
        self.conn.commit()

    def setup_settings_table(self):
        self.cursor.execute("CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT)")
        default_settings = {
            'association_name': 'الجمعية الدينية لمسجد سعد إبن أبي الوقاص',
            'address': 'حي 400 مسكن واد نشو - غرداية',
            'phone': '0662498730'
        }
        for key, value in default_settings.items():
            self.cursor.execute("INSERT OR IGNORE INTO settings (key, value) VALUES (?, ?)", (key, value))
        self.conn.commit()

    def get_all_settings(self):
        self.cursor.execute("SELECT key, value FROM settings")
        return {row['key']: row['value'] for row in self.cursor.fetchall()}

    def update_setting(self, key, value):
        self.cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", (key, value))
        self.conn.commit()

    def init_default_categories(self):
        default_incomes = ["تبرعات أفراد", "تبرعات مؤسسات", "تبرعات عينية", "منح", "اشتراكات", "بيع أصول", "أخرى"]
        default_expenses = ["مواد بناء", "صيانة وإصلاحات", "أجور عمال", "فواتير (ماء، كهرباء)", "مواد تنظيف", "ضيافة", "أخرى"]
        self.cursor.execute("SELECT COUNT(*) FROM income_categories")
        if self.cursor.fetchone()[0] == 0:
            for cat in default_incomes: self.cursor.execute("INSERT OR IGNORE INTO income_categories (name) VALUES (?)", (cat,))
        self.cursor.execute("SELECT COUNT(*) FROM expense_categories")
        if self.cursor.fetchone()[0] == 0:
            for cat in default_expenses: self.cursor.execute("INSERT OR IGNORE INTO expense_categories (name) VALUES (?)", (cat,))
        self.conn.commit()

    def get_categories(self, table_name):
        self.cursor.execute(f"SELECT id, name FROM {table_name} ORDER BY name")
        return self.cursor.fetchall()

    def add_category(self, table_name, name):
        try:
            self.cursor.execute(f"INSERT INTO {table_name} (name) VALUES (?)", (name,))
            self.conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False

    def update_category(self, table_name, cat_id, new_name):
        try:
            self.cursor.execute(f"UPDATE {table_name} SET name = ? WHERE id = ?", (new_name, cat_id))
            self.conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False

    def delete_category(self, table_name, cat_id):
        self.cursor.execute(f"DELETE FROM {table_name} WHERE id=?", (cat_id,))
        self.conn.commit()

    def add_user(self, username, password, role, must_change_password=0):
        hashed_password = self.hash_password(password)
        try:
            self.cursor.execute("INSERT INTO users (username, password, role, must_change_password) VALUES (?, ?, ?, ?)", (username, hashed_password, role, must_change_password))
            self.conn.commit()
            return True
        except sqlite3.IntegrityError: return False

    def get_user(self, username):
        self.cursor.execute("SELECT id, username, password, role, must_change_password FROM users WHERE username=?", (username,))
        return self.cursor.fetchone()

    def get_all_users(self):
        self.cursor.execute("SELECT id, username, role FROM users ORDER BY id")
        return self.cursor.fetchall()

    def update_user(self, user_id, username, password, role, must_change_password=None):
        query = "UPDATE users SET username=?, role=?"
        params = [username, role]
        if password:
            query += ", password=?"
            params.append(self.hash_password(password))
        if must_change_password is not None:
            query += ", must_change_password=?"
            params.append(must_change_password)
        query += " WHERE id=?"
        params.append(user_id)
        self.cursor.execute(query, tuple(params))
        self.conn.commit()

    def delete_user(self, user_id):
        self.cursor.execute("DELETE FROM users WHERE id=?", (user_id,))
        self.conn.commit()

    def log_action(self, username, action, details=""):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.cursor.execute("INSERT INTO audit_log (timestamp, username, action, details) VALUES (?, ?, ?, ?)", (timestamp, username, action, details))
        self.conn.commit()

    def get_audit_logs(self):
        self.cursor.execute("SELECT id, timestamp, username, action, details FROM audit_log ORDER BY timestamp DESC")
        return self.cursor.fetchall()

    def add_transaction(self, table_name, data):
        if table_name == "incomes":
            query = "INSERT INTO incomes (amount, date, category, description, notes, payer, attachment_path) VALUES (?, ?, ?, ?, ?, ?, ?)"
            params = (data['amount'], data['date'], data['category'], data['description'], data['notes'], data['payer'], data['attachment_path'])
        else:
            query = "INSERT INTO expenses (amount, date, category, description, notes, attachment_path) VALUES (?, ?, ?, ?, ?, ?)"
            params = (data['amount'], data['date'], data['category'], data['description'], data['notes'], data['attachment_path'])
        self.cursor.execute(query, params)
        last_id = self.cursor.lastrowid
        self.conn.commit()
        return last_id

    def get_transaction_by_id(self, table_name, record_id):
        self.cursor.execute(f"SELECT * FROM {table_name} WHERE id = ?", (record_id,))
        return self.cursor.fetchone()

    def get_transactions(self, table_name, start_date=None, end_date=None):
        query = f"SELECT * FROM {table_name}"
        params = []
        if start_date and end_date:
            query += " WHERE date >= ? AND date <= ?"
            params.extend([start_date, end_date])
        query += " ORDER BY date DESC, id DESC"
        self.cursor.execute(query, tuple(params))
        return self.cursor.fetchall()

    def update_transaction(self, table_name, record_id, data):
        if table_name == "incomes":
            query = "UPDATE incomes SET amount=?, date=?, category=?, description=?, notes=?, payer=?, attachment_path=? WHERE id=?"
            params = (data['amount'], data['date'], data['category'], data['description'], data['notes'], data['payer'], data['attachment_path'], record_id)
        else:
            query = "UPDATE expenses SET amount=?, date=?, category=?, description=?, notes=?, attachment_path=? WHERE id=?"
            params = (data['amount'], data['date'], data['category'], data['description'], data['notes'], data['attachment_path'], record_id)
        self.cursor.execute(query, params)
        self.conn.commit()

    def delete_transaction(self, table_name, record_id):
        self.cursor.execute(f"DELETE FROM {table_name} WHERE id=?", (record_id,))
        self.conn.commit()

    def clear_all_transactions(self):
        self.cursor.execute("DELETE FROM incomes")
        self.cursor.execute("DELETE FROM expenses")
        self.conn.commit()

    def add_member(self, data):
        query = "INSERT INTO members (full_name, join_date, phone, address, status, notes) VALUES (?, ?, ?, ?, ?, ?)"
        params = (data['full_name'], data['join_date'], data['phone'], data['address'], data['status'], data['notes'])
        self.cursor.execute(query, params)
        self.conn.commit()
        return self.cursor.lastrowid

    def update_member(self, member_id, data):
        query = "UPDATE members SET full_name=?, join_date=?, phone=?, address=?, status=?, notes=? WHERE id=?"
        params = (data['full_name'], data['join_date'], data['phone'], data['address'], data['status'], data['notes'], member_id)
        self.cursor.execute(query, params)
        self.conn.commit()

    def delete_member(self, member_id):
        self.cursor.execute("DELETE FROM members WHERE id=?", (member_id,))
        self.conn.commit()

    def get_all_members(self):
        self.cursor.execute("SELECT * FROM members ORDER BY full_name")
        return self.cursor.fetchall()

    def get_member_by_id(self, member_id):
        self.cursor.execute("SELECT * FROM members WHERE id = ?", (member_id,))
        return self.cursor.fetchone()

    def add_activity(self, data):
        query = "INSERT INTO activities (name, date, location, description) VALUES (?, ?, ?, ?)"
        params = (data['name'], data['date'], data['location'], data['description'])
        self.cursor.execute(query, params)
        self.conn.commit()
        return self.cursor.lastrowid

    def update_activity(self, activity_id, data):
        query = "UPDATE activities SET name=?, date=?, location=?, description=? WHERE id=?"
        params = (data['name'], data['date'], data['location'], data['description'], activity_id)
        self.cursor.execute(query, params)
        self.conn.commit()

    def delete_activity(self, activity_id):
        self.cursor.execute("DELETE FROM activities WHERE id=?", (activity_id,))
        self.conn.commit()

    def get_all_activities(self):
        self.cursor.execute("SELECT * FROM activities ORDER BY date DESC")
        return self.cursor.fetchall()

    def get_activity_by_id(self, activity_id):
        self.cursor.execute("SELECT * FROM activities WHERE id = ?", (activity_id,))
        return self.cursor.fetchone()

    def get_attendance(self, activity_id):
        self.cursor.execute("SELECT member_id FROM activity_attendance WHERE activity_id=?", (activity_id,))
        return [row['member_id'] for row in self.cursor.fetchall()]

    def update_attendance(self, activity_id, member_ids):
        self.cursor.execute("DELETE FROM activity_attendance WHERE activity_id=?", (activity_id,))
        for member_id in member_ids:
            self.cursor.execute("INSERT INTO activity_attendance (activity_id, member_id) VALUES (?, ?)", (activity_id, member_id))
        self.conn.commit()


# =================================================================
# الفئة الرئيسية للتطبيق (App)
# =================================================================
class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        FontManager.initialize_fonts()

        # --- Setup data directories in a user-writable location ---
        self.data_dir = os.path.join(os.path.expanduser("~"), "MasjidProData")
        self.attachments_dir = os.path.join(self.data_dir, "attachments")
        self.receipts_dir = os.path.join(self.data_dir, "receipts")

        try:
            os.makedirs(self.attachments_dir, exist_ok=True)
            os.makedirs(self.receipts_dir, exist_ok=True)
        except OSError as e:
            messagebox.showerror("خطأ فادح", f"فشل في إنشاء مجلدات البيانات الضرورية:\n{e}")
            self.quit()
            return
        
        # --- Database setup ---
        db_name = "masjid_pro_database_v6.db"
        app_dir_db_path = os.path.abspath(db_name)
        data_dir_db_path = os.path.join(self.data_dir, db_name)

        # Migrate database if it exists in app folder but not in data folder
        if os.path.exists(app_dir_db_path) and not os.path.exists(data_dir_db_path):
            try:
                shutil.copy2(app_dir_db_path, data_dir_db_path)
            except Exception as e:
                print(f"WARNING: Could not copy database from app directory: {e}")

        self.db = DatabaseManager(data_dir_db_path)
        self.current_user = None
        self.current_role = None

        try:
            self.icons = IconManager.load_all_icons()
        except Exception as e:
            messagebox.showerror("خطأ في تحميل المكونات", f"قد لا تظهر الأيقونات بشكل صحيح.\nالخطأ: {e}", parent=self)
            self.icons = {}

        self.update_global_style()
        self.title(f"برنامج إدارة الجمعية - إصدار {APP_VERSION}")
        
        def get_resource_path(relative_path):
            try:
                base_path = sys._MEIPASS
            except Exception:
                base_path = os.path.abspath(".")
            return os.path.join(base_path, relative_path)

        icon_path = get_resource_path("icons/app_icon.ico")
        if os.path.exists(icon_path):
            self.iconbitmap(icon_path)
        else:
            print("Warning: app_icon.ico not found. The default icon will be used.")
            
        self.geometry("1366x768")
        self.minsize(1024, 700)
        self.withdraw()
        LoginWindow(self, self)
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def update_global_style(self):
        style = ttk.Style(self)
        style.theme_use("default")
        theme_fg_color = ctk.ThemeManager.theme["CTkFrame"]["fg_color"]
        theme_text_color = ctk.ThemeManager.theme["CTkLabel"]["text_color"]
        theme_button_color = ctk.ThemeManager.theme["CTkButton"]["fg_color"]

        is_dark = ctk.get_appearance_mode() == "Dark"
        bg_color = theme_fg_color[1] if is_dark else theme_fg_color[0]
        text_color = theme_text_color[1] if is_dark else theme_text_color[0]
        selected_color = theme_button_color[1] if is_dark else theme_button_color[0]

        style.configure("Treeview", background=bg_color, foreground=text_color, rowheight=30, fieldbackground=bg_color, font=FontManager.TABLE_BODY_FONT)
        style.map('Treeview', background=[('selected', selected_color)], foreground=[('selected', "white")])
        style.configure("Treeview.Heading", font=FontManager.TABLE_HEADER_FONT, background=bg_color, foreground=text_color, relief="flat", padding=(5, 5))
        style.map("Treeview.Heading", background=[('active', selected_color)])

        if hasattr(self, 'frames') and self.frames:
            for frame in self.frames.values():
                if isinstance(frame, (DashboardFrame, ReportsFrame)) and frame.winfo_ismapped():
                    frame.on_show()

    def handle_login(self, username, role, must_change_password):
        self.current_user, self.current_role = username, role
        self.db.log_action(username, "تسجيل دخول")
        self.deiconify()
        if must_change_password:
             ChangePasswordDialog(self, self.db.get_user(username)[0])
        self.setup_main_ui()

    def on_closing(self):
        if messagebox.askyesno("تأكيد الخروج", "هل أنت متأكد من رغبتك في الخروج من البرنامج؟", parent=self):
            if self.current_user:
                 self.db.log_action(self.current_user, "تسجيل خروج")
            self.quit()

    def setup_main_ui(self):
        for widget in self.winfo_children(): widget.destroy()
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.sidebar_frame = SidebarFrame(self, controller=self)
        self.sidebar_frame.grid(row=0, column=0, sticky="ns")

        self.container = ctk.CTkFrame(self, corner_radius=0)
        self.container.grid(row=0, column=1, sticky="nsew")
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)

        self.frames = {}
        frames_to_load = {
            "DashboardFrame": DashboardFrame, 
            "IncomeFrame": IncomeFrame, 
            "ExpenseFrame": ExpenseFrame, 
            "MemberManagementFrame": MemberManagementFrame,
            "ActivityManagementFrame": ActivityManagementFrame,
            "ReportsFrame": ReportsFrame
        }
        if self.current_role == "مدير":
            frames_to_load.update({
                "DataManagementFrame": DataManagementFrame, 
                "UserManagementFrame": UserManagementFrame, 
                "AuditLogFrame": AuditLogFrame
            })

        for name, F in frames_to_load.items():
            try:
                frame = F(parent=self.container, controller=self)
                self.frames[name] = frame
                frame.grid(row=0, column=0, sticky="nsew")
            except Exception as e:
                messagebox.showerror("خطأ في تهيئة الواجهة", f"فشل تحميل الصفحة: {name}\n\n{traceback.format_exc()}", parent=self)

        self.show_frame("DashboardFrame")

    def show_frame(self, page_name):
        try:
            frame = self.frames[page_name]
            if hasattr(frame, "on_show"):
                frame.on_show()
            frame.tkraise()
        except Exception as e:
            messagebox.showerror("خطأ فادح", f"حدث خطأ أثناء عرض صفحة: {page_name}\n\n{traceback.format_exc()}", parent=self)

    def refresh_all_data(self):
        for frame in self.frames.values():
            if hasattr(frame, "on_show"):
                frame.on_show()

    def export_full_backup(self):
        filepath = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")], title="حفظ النسخة الاحتياطية الكاملة", initialfile=f"نسخة_احتياطية_كاملة_{datetime.now().strftime('%Y-%m-%d')}.xlsx")
        if filepath:
            self.generate_excel_report(filepath)

    def generate_excel_report(self, filepath, start_date=None, end_date=None):
        try:
            incomes_data = self.db.get_transactions('incomes', start_date, end_date)
            expenses_data = self.db.get_transactions('expenses', start_date, end_date)
            df_incomes = pd.DataFrame([dict(row) for row in incomes_data])
            df_expenses = pd.DataFrame([dict(row) for row in expenses_data])

            total_income = df_incomes['amount'].sum() if not df_incomes.empty else 0
            total_expense = df_expenses['amount'].sum() if not df_expenses.empty else 0
            balance = total_income - total_expense

            df_incomes_export = df_incomes[['date', 'payer', 'description', 'category', 'amount']].copy() if not df_incomes.empty else pd.DataFrame(columns=["التاريخ", "الدافع/المصدر", "الوصف", "الفئة", "المبلغ"])
            if not df_incomes.empty:
                df_incomes_export.columns = ["التاريخ", "الدافع/المصدر", "الوصف", "الفئة", "المبلغ"]
                df_incomes_export['المبلغ'] = df_incomes_export['المبلغ'].apply(lambda x: f"{x:,.2f} د.ج")

            df_expenses_export = df_expenses[['date', 'description', 'category', 'amount']].copy() if not df_expenses.empty else pd.DataFrame(columns=["التاريخ", "الوصف", "الفئة", "المبلغ"])
            if not df_expenses.empty:
                df_expenses_export.columns = ["التاريخ", "الوصف", "الفئة", "المبلغ"]
                df_expenses_export['المبلغ'] = df_expenses_export['المبلغ'].apply(lambda x: f"{x:,.2f} د.ج")

            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
                title_font_excel = OpenpyxlFont(name='Calibri', size=18, bold=True, color='FFFFFF')
                title_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
                center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
                right_align = Alignment(horizontal='right', vertical='center')
                header_font = OpenpyxlFont(name='Calibri', size=12, bold=True, color='FFFFFF')
                header_fill = PatternFill(start_color="808080", end_color="808080", fill_type="solid")
                cell_font = OpenpyxlFont(name='Calibri', size=11)

                ws_summary = writer.book.create_sheet(title="ملخص مالي", index=0)
                ws_summary.sheet_view.rightToLeft = True
                report_period = f"من {start_date} إلى {end_date}" if start_date and end_date else "لكامل الفترة"
                ws_summary.merge_cells('B2:E2'); title_cell = ws_summary['B2']
                title_cell.value = f"التقرير المالي {report_period}"; title_cell.font = title_font_excel; title_cell.fill = title_fill; title_cell.alignment = center_align; ws_summary.row_dimensions[2].height = 25

                summary_data = { "إجمالي المداخيل": total_income, "إجمالي المصروفات": total_expense, "الرصيد الصافي": balance }
                label_font = OpenpyxlFont(name='Calibri', size=12, bold=True)
                value_font_green = OpenpyxlFont(name='Calibri', size=12, bold=True, color='00B050')
                value_font_red = OpenpyxlFont(name='Calibri', size=12, bold=True, color='C00000')
                value_font_blue = OpenpyxlFont(name='Calibri', size=12, bold=True, color='0070C0')
                row_num = 5
                ws_summary.merge_cells(f'C{row_num}:D{row_num}'); ws_summary.merge_cells(f'C{row_num+1}:D{row_num+1}'); ws_summary.merge_cells(f'C{row_num+2}:D{row_num+2}')
                for label, value in summary_data.items():
                    label_cell, value_cell = ws_summary[f'E{row_num}'], ws_summary[f'C{row_num}']
                    label_cell.value = label; value_cell.value = f"{value:,.2f} د.ج"
                    label_cell.font = label_font; label_cell.alignment = right_align; label_cell.border = thin_border
                    value_cell.alignment = center_align; value_cell.border = thin_border; ws_summary[f'D{row_num}'].border = thin_border
                    if "المداخيل" in label: value_cell.font = value_font_green
                    elif "المصاريف" in label: value_cell.font = value_font_red
                    else: value_cell.font = value_font_blue
                    row_num += 1
                ws_summary.column_dimensions['C'].width = 15; ws_summary.column_dimensions['D'].width = 15; ws_summary.column_dimensions['E'].width = 25

                df_incomes_export.to_excel(writer, sheet_name='المداخيل', index=False, startrow=2)
                df_expenses_export.to_excel(writer, sheet_name='المصاريف', index=False, startrow=2)

                for sheet_name, title_text, df in [('المداخيل', 'قائمة المداخيل', df_incomes_export), ('المصاريف', 'قائمة المصاريف', df_expenses_export)]:
                    ws = writer.book[sheet_name]; ws.sheet_view.rightToLeft = True
                    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=df.shape[1] if not df.empty else 1)
                    main_title_cell = ws.cell(row=1, column=1); main_title_cell.value = f"{title_text} {report_period}"; main_title_cell.font = title_font_excel; main_title_cell.fill = title_fill; main_title_cell.alignment = center_align; ws.row_dimensions[1].height = 22
                    for c_idx, col_name in enumerate(df.columns, 1):
                        header_cell = ws.cell(row=3, column=c_idx); header_cell.value = col_name; header_cell.font = header_font; header_cell.fill = header_fill; header_cell.border = thin_border; header_cell.alignment = center_align
                    for r_idx in range(4, ws.max_row + 1):
                        for c_idx in range(1, ws.max_column + 1):
                            cell = ws.cell(row=r_idx, column=c_idx); cell.font = cell_font; cell.border = thin_border; cell.alignment = right_align
                    for col in ws.columns:
                        max_length = 0; column_letter = get_column_letter(col[0].column)
                        for cell in col:
                            if cell.value and len(str(cell.value)) > max_length: max_length = len(str(cell.value))
                        adjusted_width = (max_length + 2) * 1.2
                        ws.column_dimensions[column_letter].width = adjusted_width if adjusted_width > 12 else 12

            self.db.log_action(self.current_user, "تصدير تقرير Excel", f"الملف: {os.path.basename(filepath)}")
            messagebox.showinfo("نجاح", "تم إنشاء تقرير Excel بنجاح!", parent=self)
        except Exception as e:
            messagebox.showerror("خطأ", f"فشل إنشاء التقرير:\n{e}", parent=self)
            traceback.print_exc()

    def restore_from_excel(self):
        if not messagebox.askyesno("تأكيد الاستعادة", "هل أنت متأكد من استعادة البيانات؟\nسيتم حذف جميع البيانات الحالية واستبدالها ببيانات النسخة الاحتياطية.\nلا يمكن التراجع عن هذا الإجراء.", icon='warning', parent=self):
            return
        filepath = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")], title="اختيار ملف النسخة الاحتياطية")
        if not filepath: return

        try:
            df_incomes = pd.read_excel(filepath, sheet_name='المداخيل', skiprows=2)
            df_expenses = pd.read_excel(filepath, sheet_name='المصاريف', skiprows=2)

            def clean_currency(value):
                if isinstance(value, str): return float(value.replace('د.ج', '').replace(',', '').strip())
                return float(value)

            df_incomes['المبلغ'] = df_incomes['المبلغ'].apply(clean_currency)
            df_expenses['المبلغ'] = df_expenses['المبلغ'].apply(clean_currency)

            required_income_cols, required_expense_cols = ["التاريخ", "الدافع/المصدر", "الوصف", "الفئة", "المبلغ"], ["التاريخ", "الوصف", "الفئة", "المبلغ"]
            if not all(col in df_incomes.columns for col in required_income_cols) or not all(col in df_expenses.columns for col in required_expense_cols):
                 messagebox.showerror("خطأ في الملف", "ملف النسخة الاحتياطية غير صالح أو تالف.", parent=self); return

            self.db.clear_all_transactions()
            if os.path.exists(self.controller.attachments_dir): shutil.rmtree(self.controller.attachments_dir)
            os.makedirs(self.controller.attachments_dir)

            for _, row in df_incomes.iterrows():
                self.db.add_transaction('incomes', {'amount': row['المبلغ'], 'date': str(row['التاريخ']).split(" ")[0], 'category': row['الفئة'], 'description': row['الوصف'], 'payer': row['الدافع/المصدر'], 'notes': '', 'attachment_path': None})
            for _, row in df_expenses.iterrows():
                self.db.add_transaction('expenses', {'amount': row['المبلغ'], 'date': str(row['التاريخ']).split(" ")[0], 'category': row['الفئة'], 'description': row['الوصف'], 'notes': '', 'attachment_path': None})

            self.db.log_action(self.current_user, "استعادة نسخة احتياطية", f"الملف: {os.path.basename(filepath)}")
            messagebox.showinfo("نجاح", "تم استعادة البيانات بنجاح.", parent=self)
            self.refresh_all_data()
        except Exception as e:
            messagebox.showerror("خطأ", f"فشل استعادة البيانات:\n{e}", parent=self)
            traceback.print_exc()

# =================================================================
# الشريط الجانبي (SidebarFrame)
# =================================================================
class SidebarFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, corner_radius=0)
        self.controller = controller

        # Layout: 0=Top Labels, 1=Scrollable Buttons (expands), 2=Bottom Buttons (fixed)
        self.grid_rowconfigure(1, weight=1)
        
        # --- Top Labels Frame ---
        top_frame = ctk.CTkFrame(self, fg_color="transparent")
        top_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")
        ctk.CTkLabel(top_frame, text="جمعية المسجد", font=FontManager.H1_FONT).pack()
        ctk.CTkLabel(top_frame, text=f"مرحباً، {self.controller.current_user}\n({self.controller.current_role})", 
                     font=FontManager.APP_FONT, justify="right").pack(pady=(0, 10))

        # --- Scrollable Frame for Navigation Buttons ---
        scrollable_buttons_frame = ctk.CTkScrollableFrame(self, label_text="", fg_color="transparent")
        scrollable_buttons_frame.grid(row=1, column=0, padx=15, pady=0, sticky="nsew")

        buttons = {
            "لوحة التحكم": "DashboardFrame", 
            "إدارة المداخيل": "IncomeFrame", 
            "إدارة المصروفات": "ExpenseFrame",
            "إدارة الأعضاء": "MemberManagementFrame",
            "إدارة الأنشطة": "ActivityManagementFrame",
            "التقارير": "ReportsFrame"
        }
        button_icons = {
            "لوحة التحكم": "dashboard", 
            "إدارة المداخيل": "income", 
            "إدارة المصروفات": "expense",
            "إدارة الأعضاء": "users",
            "إدارة الأنشطة": "activity",
            "التقارير": "reports"
        }

        for text, frame_name in buttons.items():
            ctk.CTkButton(scrollable_buttons_frame, text=text, image=self.controller.icons.get(button_icons[text]), 
                          compound="right", anchor="e", command=lambda f=frame_name: self.controller.show_frame(f), 
                          font=FontManager.BUTTON_FONT).pack(fill="x", padx=5, pady=5)
        
        if self.controller.current_role == "مدير":
            ctk.CTkLabel(scrollable_buttons_frame, text="أدوات المدير", 
                         font=ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=14, weight="bold"), 
                         anchor="e").pack(fill="x", padx=5, pady=(15, 5))
            
            admin_buttons = {"إدارة البيانات": "DataManagementFrame", "إدارة المستخدمين": "UserManagementFrame", "سجل التدقيق": "AuditLogFrame"}
            admin_icons = {"إدارة البيانات": "data", "إدارة المستخدمين": "users", "سجل التدقيق": "audit"}
            for text, frame_name in admin_buttons.items():
                ctk.CTkButton(scrollable_buttons_frame, text=text, image=self.controller.icons.get(admin_icons[text]), 
                              compound="right", anchor="e", command=lambda f=frame_name: self.controller.show_frame(f), 
                              font=FontManager.BUTTON_FONT).pack(fill="x", padx=5, pady=5)
        
        # --- Bottom Fixed Frame ---
        bottom_frame = ctk.CTkFrame(self, fg_color="transparent")
        bottom_frame.grid(row=2, column=0, padx=20, pady=10, sticky="sew")

        self.appearance_mode_frame = ctk.CTkFrame(bottom_frame, fg_color="transparent")
        self.appearance_mode_frame.pack(fill="x", pady=5)
        ctk.CTkLabel(self.appearance_mode_frame, text="المظهر:", font=FontManager.BUTTON_FONT, anchor="w").pack(side="right", padx=5)
        self.appearance_mode_switch = ctk.CTkSwitch(self.appearance_mode_frame, text="", command=self.toggle_appearance_mode, onvalue="dark", offvalue="light")
        self.appearance_mode_switch.pack(side="left")
        self.appearance_mode_switch.select() if ctk.get_appearance_mode().lower() == "dark" else self.appearance_mode_switch.deselect()
        
        ctk.CTkButton(bottom_frame, text="حول البرنامج", image=self.controller.icons.get("about"), compound="right", command=self.show_about_window, font=FontManager.BUTTON_FONT).pack(fill="x", pady=5)
        ctk.CTkButton(bottom_frame, text="تسجيل الخروج", image=self.controller.icons.get("logout"), compound="right", command=parent.on_closing, fg_color="transparent", border_width=2, text_color=("gray10", "#DCE4EE"), font=FontManager.BUTTON_FONT).pack(fill="x", pady=(5,10))

    def toggle_appearance_mode(self):
        new_mode = self.appearance_mode_switch.get()
        ctk.set_appearance_mode(new_mode)
        self.controller.update_global_style()

    def show_about_window(self):
        AboutWindow(self)

# =================================================================
# باقي الفئات والواجهات
# =================================================================
class LoginWindow(ctk.CTkToplevel):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.title("تسجيل الدخول"); self.geometry("450x350"); self.resizable(False, False); self.transient(parent)
        self.protocol("WM_DELETE_WINDOW", self.controller.quit)
        frame = ctk.CTkFrame(self, fg_color="transparent"); frame.pack(expand=True, padx=20, pady=20)
        ctk.CTkLabel(frame, text="برنامج إدارة الجمعية", font=FontManager.TITLE_FONT).pack(pady=(0, 20))
        ctk.CTkLabel(frame, text="اسم المستخدم", font=FontManager.APP_FONT).pack(pady=(10, 5))
        self.username_entry = ctk.CTkEntry(frame, width=300, justify='center', font=FontManager.INPUT_FONT); self.username_entry.pack()
        ctk.CTkLabel(frame, text="كلمة المرور", font=FontManager.APP_FONT).pack(pady=(10, 5))
        self.password_entry = ctk.CTkEntry(frame, width=300, show="*", justify='center', font=FontManager.INPUT_FONT); self.password_entry.pack()
        self.password_entry.bind("<Return>", self.login); self.username_entry.focus()
        ctk.CTkButton(frame, text="دخول", command=self.login, font=FontManager.BUTTON_FONT, width=300).pack(pady=20)
        self.grab_set()

    def login(self, event=None):
        username = self.username_entry.get(); password = self.password_entry.get()
        if not username or not password: messagebox.showerror("خطأ", "يرجى إدخال اسم المستخدم وكلمة المرور.", parent=self); return
        user_data = self.controller.db.get_user(username)
        if user_data and self.controller.db.hash_password(password) == user_data['password']:
            self.destroy(); self.controller.handle_login(username, user_data['role'], user_data['must_change_password'])
        else: messagebox.showerror("خطأ", "اسم المستخدم أو كلمة المرور غير صحيحة.", parent=self)

class BaseDataFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller, self.db, self.is_admin = controller, controller.db, controller.current_role == "مدير"
        self.grid_columnconfigure(0, weight=1); self.grid_rowconfigure(1, weight=1)

    def setup_treeview(self, columns, headings):
        tree_frame = ctk.CTkFrame(self); tree_frame.grid(row=1, column=0, padx=20, pady=(0, 20), sticky="nsew")
        tree_frame.grid_rowconfigure(0, weight=1); tree_frame.grid_columnconfigure(0, weight=1)
        self.tree = ttk.Treeview(tree_frame, columns=columns, show="headings", selectmode="browse")
        for col, head in zip(columns, headings):
            self.tree.heading(col, text=head, anchor='center'); self.tree.column(col, anchor='e', width=120)
        self.tree.grid(row=0, column=0, sticky="nsew")
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set); scrollbar.grid(row=0, column=1, sticky="ns")
        return self.tree

class TransactionFrame(BaseDataFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        self.all_data = []
        self.setup_ui()

    def setup_ui(self):
        top_frame = ctk.CTkFrame(self); top_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        ctk.CTkLabel(top_frame, text=self.get_title(), font=FontManager.TITLE_FONT).pack(side="right", padx=10, pady=5)
        self.button_panel = ctk.CTkFrame(top_frame, fg_color="transparent"); self.button_panel.pack(side="left", padx=5, pady=10)
        self.attachment_button = ctk.CTkButton(self.button_panel, text="عرض المرفق", image=self.controller.icons.get("attachment"), compound="right", command=self.view_attachment, font=FontManager.BUTTON_FONT, state="disabled")
        self.attachment_button.pack(side="left", padx=5)
        ctk.CTkButton(self.button_panel, text="جديد", image=self.controller.icons.get("add"), compound="right", command=self.add_new_item, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        if self.is_admin:
            ctk.CTkButton(self.button_panel, text="تعديل", image=self.controller.icons.get("edit"), compound="right", command=self.edit_selected_item, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
            ctk.CTkButton(self.button_panel, text="حذف", image=self.controller.icons.get("delete"), compound="right", command=self.delete_selected_item, fg_color="#D2042D", hover_color="#990000", font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        search_frame = ctk.CTkFrame(top_frame, fg_color="transparent"); search_frame.pack(side="left", padx=20, pady=10, expand=True, fill="x")
        self.search_entry = ctk.CTkEntry(search_frame, placeholder_text="ابحث هنا...", font=FontManager.INPUT_FONT, width=300); self.search_entry.pack(side="right", fill="x", expand=True)
        self.search_entry.bind("<KeyRelease>", self.filter_table)
        ctk.CTkLabel(search_frame, text="", image=self.controller.icons.get("search")).pack(side="right", padx=5)
        self.tree = self.setup_treeview(self.get_columns(), self.get_headings()); self.configure_tree_columns(); self.tree.bind("<<TreeviewSelect>>", self.on_item_select)

    def configure_tree_columns(self):
        self.tree.column("id", width=60, anchor='center'); self.tree.column("amount", width=150, anchor='center'); self.tree.column("description", width=350); self.tree.column("attachment", width=80, anchor='center')

    def on_item_select(self, event=None):
        selected_item = self.tree.focus()
        if not selected_item:
            self.attachment_button.configure(state="disabled")
            return
        item_values = self.tree.item(selected_item)["values"]
        if not item_values:
            self.attachment_button.configure(state="disabled")
            return
        item_id = self.get_id_from_tree_values(item_values)
        full_data = self.db.get_transaction_by_id(self.table_name, item_id)
        attachment_path = full_data['attachment_path'] if full_data else None
        self.attachment_button.configure(state="normal" if attachment_path and os.path.exists(attachment_path) else "disabled")

    def view_attachment(self):
        selected_id = self.get_selected_id()
        if not selected_id: return
        full_data = self.db.get_transaction_by_id(self.table_name, selected_id)
        attachment_path = full_data['attachment_path'] if full_data else None
        if attachment_path and os.path.exists(attachment_path): AttachmentViewer(self, attachment_path)
        else: messagebox.showwarning("تنبيه", "لا يوجد مرفق لهذه المعاملة أو أن الملف مفقود.", parent=self)

    def on_show(self):
        self.fetch_data(); self.populate_table(self.all_data); self.search_entry.delete(0, 'end'); self.attachment_button.configure(state="disabled")

    def fetch_data(self):
        self.all_data = self.db.get_transactions(self.table_name)

    def filter_table(self, event=None):
        search_term = self.search_entry.get().lower()
        if not search_term:
            self.populate_table(self.all_data)
            return
        filtered_data = [dict(row) for row in self.all_data if any(search_term in str(field).lower() for field in dict(row).values())]
        self.populate_table(filtered_data)

    def get_id_from_tree_values(self, values):
        return values[-1]

    def get_selected_id(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("تنبيه", "يرجى تحديد عنصر من الجدول أولاً.", parent=self)
            return None
        item_values = self.tree.item(selected_item)["values"]
        if not item_values:
            return None
        return self.get_id_from_tree_values(item_values)

class IncomeFrame(TransactionFrame):
    def __init__(self, parent, controller):
        self.table_name = 'incomes'
        super().__init__(parent, controller)
        self.print_button = ctk.CTkButton(self.button_panel, text="إنشاء وصل (Word)", image=self.controller.icons.get("print"), compound="right", command=self.print_receipt, font=FontManager.BUTTON_FONT, state="disabled")
        self.print_button.pack(side="left", padx=5)

    def on_item_select(self, event=None):
        super().on_item_select(event)
        self.print_button.configure(state="normal" if self.tree.focus() else "disabled")

    def get_title(self): return "إدارة المداخيل"
    def get_columns(self): return ("attachment", "description", "payer", "category", "date", "amount", "id")
    def get_headings(self): return ("مرفق", "الوصف", "الدافع", "الفئة", "التاريخ", "المبلغ", "المعرف")
    def configure_tree_columns(self): super().configure_tree_columns(); self.tree.column("payer", width=200)

    def populate_table(self, data):
        self.tree.delete(*self.tree.get_children())
        for row in data:
            inc = dict(row)
            has_attachment = "نعم" if inc.get('attachment_path') and os.path.exists(inc['attachment_path']) else "لا"
            amount_str = f"{inc.get('amount', 0):,.2f}\u200e د.ج"
            values = (has_attachment, inc.get('description', ''), inc.get('payer', ''), inc.get('category', ''), inc.get('date', ''), amount_str, inc.get('id', ''))
            self.tree.insert("", "end", values=values)

    def add_new_item(self):
        dialog = DataEntryDialog(self, title="إضافة مدخول جديد", fields=self.get_income_fields(), db=self.db, table_type="income")
        data = dialog.get_data()
        if data:
            new_id = self.db.add_transaction("incomes", data)
            self.db.log_action(self.controller.current_user, "إضافة مدخول", f"المبلغ: {data['amount']}")
            self.on_show()
            if messagebox.askyesno("إنشاء وصل", "تم حفظ المدخول بنجاح. هل تريد إنشاء وصل الآن؟", parent=self):
                income_data = self.db.get_transaction_by_id('incomes', new_id)
                if income_data:
                    self.print_receipt_for_data(income_data)

    def edit_selected_item(self):
        income_id = self.get_selected_id()
        if not income_id: return
        full_data = self.db.get_transaction_by_id(self.table_name, income_id)
        if not full_data: return
        dialog = DataEntryDialog(self, title="تعديل مدخول", fields=self.get_income_fields(), initial_data=dict(full_data), db=self.db, table_type="income")
        updated_data = dialog.get_data()
        if updated_data:
            self.db.update_transaction("incomes", income_id, updated_data)
            self.db.log_action(self.controller.current_user, "تعديل مدخول", f"معرف: {income_id}")
            self.on_show()

    def delete_selected_item(self):
        income_id = self.get_selected_id()
        if not income_id: return
        if messagebox.askyesno("تأكيد الحذف", "هل أنت متأكد من حذف هذا المدخول؟", icon='warning', parent=self):
            full_data = self.db.get_transaction_by_id(self.table_name, income_id)
            if full_data and full_data.get('attachment_path') and os.path.exists(full_data['attachment_path']):
                try:
                    os.remove(full_data['attachment_path'])
                except OSError as e:
                    print(f"Error deleting attachment: {e}")
            self.db.delete_transaction("incomes", income_id)
            self.db.log_action(self.controller.current_user, "حذف مدخول", f"معرف: {income_id}")
            self.on_show()

    def print_receipt(self):
        income_id = self.get_selected_id()
        if not income_id: return
        income_data = self.db.get_transaction_by_id('incomes', income_id)
        if income_data:
            self.print_receipt_for_data(income_data)

    def print_receipt_for_data(self, data):
        receipt_id = data['id'] if data else 'N/A'
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="حفظ وصل Word",
            initialfile=f"وصل_رقم_{receipt_id}.docx"
        )
        if not filepath:
            return

        try:
            settings = self.controller.db.get_all_settings()
            generator = WordReceiptGenerator(dict(data), settings)
            generator.save(filepath)

            messagebox.showinfo("نجاح", f"تم حفظ الوصل في:\n{filepath}\nسيتم الآن فتح الملف.", parent=self)

            if sys.platform == 'win32':
                os.startfile(filepath)
            elif sys.platform == 'darwin':
                subprocess.run(['open', filepath], check=True)
            else:
                subprocess.run(['xdg-open', filepath], check=True)
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("خطأ", f"فشل إنشاء أو فتح ملف الوصل:\n{e}", parent=self)


    def get_income_fields(self): return {"amount": {"label": "المبلغ (د.ج)", "type": "number", "required": True}, "date": {"label": "التاريخ", "type": "date", "required": True}, "category": {"label": "الفئة", "type": "combo", "required": True}, "payer": {"label": "اسم الدافع/المتبرع", "type": "text", "required": True}, "description": {"label": "الوصف", "type": "text"}, "notes": {"label": "ملاحظات", "type": "textarea"}}

class ExpenseFrame(TransactionFrame):
    def __init__(self, parent, controller):
        self.table_name = 'expenses'
        super().__init__(parent, controller)

    def get_title(self): return "إدارة المصروفات"
    def get_columns(self): return ("attachment", "description", "category", "date", "amount", "id")
    def get_headings(self): return ("مرفق", "الوصف", "الفئة", "التاريخ", "المبلغ", "المعرف")

    def populate_table(self, data):
        self.tree.delete(*self.tree.get_children())
        for row in data:
            exp = dict(row)
            has_attachment = "نعم" if exp.get('attachment_path') and os.path.exists(exp['attachment_path']) else "لا"
            amount_str = f"{exp.get('amount', 0):,.2f}\u200e د.ج"
            values = (has_attachment, exp.get('description',''), exp.get('category',''), exp.get('date',''), amount_str, exp.get('id',''))
            self.tree.insert("", "end", values=values)

    def add_new_item(self):
        dialog = DataEntryDialog(self, title="إضافة مصروف جديد", fields=self.get_expense_fields(), db=self.db, table_type="expense"); data = dialog.get_data()
        if data:
            self.db.add_transaction("expenses", data)
            self.db.log_action(self.controller.current_user, "إضافة مصروف", f"المبلغ: {data['amount']}")
            self.on_show()

    def edit_selected_item(self):
        expense_id = self.get_selected_id()
        if not expense_id: return
        full_data = self.db.get_transaction_by_id(self.table_name, expense_id)
        if not full_data: return
        dialog = DataEntryDialog(self, title="تعديل مصروف", fields=self.get_expense_fields(), initial_data=dict(full_data), db=self.db, table_type="expense")
        updated_data = dialog.get_data()
        if updated_data:
            self.db.update_transaction("expenses", expense_id, updated_data)
            self.db.log_action(self.controller.current_user, "تعديل مصروف", f"معرف: {expense_id}")
            self.on_show()

    def delete_selected_item(self):
        expense_id = self.get_selected_id()
        if not expense_id: return
        if messagebox.askyesno("تأكيد الحذف", "هل أنت متأكد من حذف هذا المصروف؟", icon='warning', parent=self):
            full_data = self.db.get_transaction_by_id(self.table_name, expense_id)
            if full_data and full_data.get('attachment_path') and os.path.exists(full_data['attachment_path']):
                try:
                    os.remove(full_data['attachment_path'])
                except OSError as e:
                    print(f"Error deleting attachment: {e}")
            self.db.delete_transaction("expenses", expense_id)
            self.db.log_action(self.controller.current_user, "حذف مصروف", f"معرف: {expense_id}")
            self.on_show()

    def get_expense_fields(self): return {"amount": {"label": "المبلغ (د.ج)", "type": "number", "required": True}, "date": {"label": "التاريخ", "type": "date", "required": True}, "category": {"label": "الفئة", "type": "combo", "required": True}, "description": {"label": "الوصف", "type": "text"}, "notes": {"label": "ملاحظات", "type": "textarea"}}

class DataEntryDialog(ctk.CTkToplevel):
    def __init__(self, parent, title, fields, db, table_type, initial_data=None):
        super().__init__(parent)
        self.transient(parent)
        self.title(title)
        self.geometry("700x650")
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", self.cancel)

        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.fields, self.db, self.table_type, self.entries, self.data = fields, db, table_type, {}, None
        self.attachment_path = initial_data.get('attachment_path') if initial_data else None
        self.thumbnail_image = None

        main_frame = ctk.CTkScrollableFrame(self)
        main_frame.grid(row=0, column=0, padx=10, pady=(10,0), sticky="nsew")
        main_frame.grid_columnconfigure(1, weight=1)

        for i, (key, props) in enumerate(self.fields.items()):
            ctk.CTkLabel(main_frame, text=props['label'] + (" *" if props.get("required") else ""), anchor="e", font=FontManager.APP_FONT).grid(row=i, column=0, padx=10, pady=10, sticky="w")
            if props['type'] == "textarea":
                entry = ctk.CTkTextbox(main_frame, height=100, font=FontManager.INPUT_FONT)
            elif props['type'] == 'combo':
                categories = [cat['name'] for cat in self.db.get_categories("income_categories" if self.table_type == "income" else "expense_categories")]
                entry = ctk.CTkComboBox(main_frame, values=categories, justify='right', font=FontManager.INPUT_FONT, dropdown_font=FontManager.INPUT_FONT, state="readonly", button_color="#3a7ebf")
                if not initial_data and categories:
                    entry.set(categories[0])
            else:
                entry = ctk.CTkEntry(main_frame, justify='right', font=FontManager.INPUT_FONT)

            if props['type'] == 'date' and not initial_data:
                entry.insert(0, datetime.now().strftime("%Y-%m-%d"))

            entry.grid(row=i, column=1, padx=10, pady=10, sticky="ew")
            self.entries[key] = entry

        i = len(self.fields)
        ctk.CTkLabel(main_frame, text="المرفق (صورة)", anchor="e", font=FontManager.APP_FONT).grid(row=i, column=0, padx=10, pady=10, sticky="w")
        attachment_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        attachment_frame.grid(row=i, column=1, padx=10, pady=10, sticky="ew")
        ctk.CTkButton(attachment_frame, text="اختر ملف...", command=self.attach_file, font=FontManager.BUTTON_FONT).pack(side="right", padx=5)
        self.attachment_label = ctk.CTkLabel(attachment_frame, text="لم يتم إرفاق ملف", anchor="w", font=FontManager.APP_FONT)
        self.attachment_label.pack(side="right", padx=5, expand=True, fill="x")

        self.thumbnail_label = ctk.CTkLabel(main_frame, text="")
        self.thumbnail_label.grid(row=i+1, column=1, pady=5, sticky="w")

        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.grid(row=1, column=0, pady=10, sticky="ew")
        button_frame.grid_columnconfigure((0, 1), weight=1) 

        save_btn = ctk.CTkButton(button_frame, text="حفظ", command=self.save, font=FontManager.BUTTON_FONT)
        save_btn.grid(row=0, column=0, padx=10, sticky="e")

        cancel_btn = ctk.CTkButton(button_frame, text="إلغاء", command=self.cancel, fg_color="gray", font=FontManager.BUTTON_FONT)
        cancel_btn.grid(row=0, column=1, padx=10, sticky="w")

        if initial_data:
            self.fill_form(initial_data)

        self.grab_set()
        self.wait_window()

    def attach_file(self):
        filepath = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.bmp"), ("All Files", "*.*")])
        if filepath:
            try:
                filename = os.path.basename(filepath)
                dest_path = os.path.join(self.master.controller.attachments_dir, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}")
                shutil.copy(filepath, dest_path)
                self.attachment_path = dest_path
                self.attachment_label.configure(text=filename)
                self.show_thumbnail()
            except Exception as e:
                messagebox.showerror("خطأ", f"فشل نسخ الملف: {e}", parent=self)
                self.attachment_path = None

    def show_thumbnail(self):
        if self.attachment_path and os.path.exists(self.attachment_path):
            try:
                img = Image.open(self.attachment_path); img.thumbnail((150, 150))
                self.thumbnail_image = ctk.CTkImage(light_image=img, dark_image=img, size=(img.width, img.height))
                self.thumbnail_label.configure(image=self.thumbnail_image, text="")
            except Exception:
                self.thumbnail_label.configure(image=None, text="لا يمكن عرض الصورة", font=FontManager.APP_FONT)
        else:
            self.thumbnail_label.configure(image=None, text="")

    def fill_form(self, initial_data):
        for key, widget in self.entries.items():
            if key in initial_data and initial_data[key] is not None:
                value = initial_data[key]
                if isinstance(widget, ctk.CTkTextbox):
                    widget.delete("1.0", "end")
                    widget.insert("1.0", value)
                elif isinstance(widget, ctk.CTkComboBox):
                    widget.set(str(value))
                else:
                    widget.delete(0, "end")
                    widget.insert(0, str(value))
        if self.attachment_path and os.path.exists(self.attachment_path):
            self.attachment_label.configure(text=os.path.basename(self.attachment_path))
            self.show_thumbnail()
        else:
            self.attachment_label.configure(text="لم يتم إرفاق ملف")
            self.show_thumbnail()

    def save(self):
        self.data = {}
        for key, props in self.fields.items():
            widget = self.entries[key]; value = widget.get("1.0", "end-1c").strip() if isinstance(widget, ctk.CTkTextbox) else widget.get().strip()
            if props.get("required") and not value:
                messagebox.showerror("خطأ", f"حقل '{props['label']}' مطلوب.", parent=self)
                return
            if props.get("type") == "number" and value:
                try: value = float(value)
                except ValueError:
                    messagebox.showerror("خطأ", "المبلغ يجب أن يكون رقماً.", parent=self)
                    return
                if value <= 0:
                    messagebox.showerror("خطأ", "المبلغ يجب أن يكون أكبر من صفر.", parent=self)
                    return
            self.data[key] = value
        self.data['attachment_path'] = self.attachment_path
        if 'payer' not in self.data: self.data['payer'] = ''
        if 'notes' not in self.data: self.data['notes'] = ''
        self.destroy()

    def cancel(self):
        self.data = None
        self.destroy()

    def get_data(self):
        return self.data

class AttachmentViewer(ctk.CTkToplevel):
    def __init__(self, parent, image_path):
        super().__init__(parent); self.title("عرض المرفق")
        self.image = None
        try:
            pil_image = Image.open(image_path); max_width, max_height = 800, 600
            if pil_image.width > max_width or pil_image.height > max_height: pil_image.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            self.geometry(f"{pil_image.width+40}x{pil_image.height+40}")
            self.image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(pil_image.width, pil_image.height))
            ctk.CTkLabel(self, image=self.image, text="").pack(padx=20, pady=20, expand=True, fill="both")
        except Exception as e:
            self.geometry("300x100"); ctk.CTkLabel(self, text=f"لا يمكن فتح الصورة.\n{e}", font=FontManager.APP_FONT).pack(pady=20, padx=20)
        self.grab_set(); self.focus()

class DashboardFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller
        self.grid_columnconfigure(0, weight=1); self.grid_rowconfigure(2, weight=1)
        title_frame = ctk.CTkFrame(self, fg_color="transparent"); title_frame.grid(row=0, column=0, columnspan=3, sticky="ew", padx=20, pady=(20, 10))
        title_frame.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(title_frame, text="لوحة التحكم الرئيسية", font=FontManager.TITLE_FONT).grid(row=0, column=0, pady=10)
        cards_frame = ctk.CTkFrame(self, fg_color="transparent"); cards_frame.grid(row=1, column=0, sticky="ew", padx=20, pady=10)
        cards_frame.grid_columnconfigure((0, 1, 2), weight=1)
        self.balance_card = self.create_summary_card(cards_frame, "الرصيد الصافي",  0, "#1E88E5")
        self.expense_card = self.create_summary_card(cards_frame, "إجمالي المصروفات", 1, "#E53935")
        self.income_card = self.create_summary_card(cards_frame, "إجمالي المداخيل", 2, "#009688")
        self.charts_frame = ctk.CTkFrame(self); self.charts_frame.grid(row=2, column=0, sticky="nsew", padx=20, pady=(0, 20))
        self.charts_frame.grid_columnconfigure((0, 1), weight=1); self.charts_frame.grid_rowconfigure(0, weight=1)
        self.monthly_chart_canvas, self.balance_chart_canvas = None, None

    def create_summary_card(self, parent, title, col, color):
        card = ctk.CTkFrame(parent, border_width=2, border_color=color, corner_radius=10)
        card.grid(row=0, column=col, padx=10, pady=10, sticky="nsew")
        card.grid_rowconfigure(1, weight=1); card.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(card, text=title, font=FontManager.H2_FONT).pack(pady=(20,10), padx=10)
        value_label = ctk.CTkLabel(card, text=f"0.00\u200e د.ج", font=ctk.CTkFont(family=FontManager.UI_FONT_FAMILY, size=32, weight="bold"), text_color=color)
        value_label.pack(pady=10, padx=20, expand=True, fill="both")
        return value_label

    def on_show(self):
        incomes = self.controller.db.get_transactions('incomes')
        expenses = self.controller.db.get_transactions('expenses')

        total_income = sum(row['amount'] for row in incomes)
        total_expense = sum(row['amount'] for row in expenses)
        balance = total_income - total_expense

        self.income_card.configure(text=f"{total_income:,.2f}\u200e د.ج")
        self.expense_card.configure(text=f"{total_expense:,.2f}\u200e د.ج")
        self.balance_card.configure(text=f"{balance:,.2f}\u200e د.ج")

        self.update_charts(incomes, expenses)

    def update_charts(self, incomes, expenses):
        if self.monthly_chart_canvas: self.monthly_chart_canvas.get_tk_widget().destroy()
        if self.balance_chart_canvas: self.balance_chart_canvas.get_tk_widget().destroy()

        df_incomes = pd.DataFrame([dict(row) for row in incomes])
        df_expenses = pd.DataFrame([dict(row) for row in expenses])

        if not df_incomes.empty or not df_expenses.empty:
            if not df_incomes.empty:
                df_incomes['date'] = pd.to_datetime(df_incomes['date'])
                income_monthly = df_incomes.set_index('date').resample('ME')['amount'].sum().rename('المداخيل')
            else:
                income_monthly = pd.Series(dtype='float64').rename('المداخيل')

            if not df_expenses.empty:
                df_expenses['date'] = pd.to_datetime(df_expenses['date'])
                expense_monthly = df_expenses.set_index('date').resample('ME')['amount'].sum().rename('المصاريف')
            else:
                expense_monthly = pd.Series(dtype='float64').rename('المصاريف')

            monthly_summary = pd.concat([income_monthly, expense_monthly], axis=1).fillna(0)

            if not monthly_summary.empty:
                monthly_summary.index = monthly_summary.index.strftime('%Y-%m')
                fig1, ax1 = self.create_styled_figure()
                monthly_summary.plot(kind='bar', ax=ax1, color=['#009688', '#E53935'])

                ax1.set_title(format_arabic('المقارنة الشهرية بين المداخيل والمصاريف'), fontproperties=FontManager.get_matplotlib_font_prop('Bold', 14))
                ax1.set_xlabel(format_arabic('الشهر'), fontproperties=FontManager.get_matplotlib_font_prop('Regular', 12))
                ax1.set_ylabel(format_arabic('المبلغ (د.ج)'), fontproperties=FontManager.get_matplotlib_font_prop('Regular', 12))

                legend_prop = FontManager.get_matplotlib_font_prop('Regular', 10)
                ax1.legend([format_arabic('المداخيل'), format_arabic('المصاريف')], prop=legend_prop)

                tick_font_prop = FontManager.get_matplotlib_font_prop('Regular', 10)
                plt.setp(ax1.get_xticklabels(), fontproperties=tick_font_prop, rotation=45, ha='right')
                plt.setp(ax1.get_yticklabels(), fontproperties=tick_font_prop)

                self.monthly_chart_canvas = FigureCanvasTkAgg(fig1, master=self.charts_frame)
                self.monthly_chart_canvas.draw()
                self.monthly_chart_canvas.get_tk_widget().grid(row=0, column=1, sticky="nsew", padx=(10, 20), pady=10)

            all_trans_list = []
            if not df_incomes.empty:
                df_incomes['amount_signed'] = df_incomes['amount']
                all_trans_list.append(df_incomes[['date', 'amount_signed']])
            if not df_expenses.empty:
                df_expenses['amount_signed'] = -df_expenses['amount']
                all_trans_list.append(df_expenses[['date', 'amount_signed']])

            if all_trans_list:
                all_trans = pd.concat(all_trans_list).sort_values('date')
                all_trans['balance'] = all_trans['amount_signed'].cumsum()
                fig2, ax2 = self.create_styled_figure()
                all_trans.plot(x='date', y='balance', ax=ax2, legend=None, color='#1E88E5')

                ax2.set_title(format_arabic('تطور الرصيد المالي'), fontproperties=FontManager.get_matplotlib_font_prop('Bold', 14))
                ax2.set_xlabel(format_arabic('التاريخ'), fontproperties=FontManager.get_matplotlib_font_prop('Regular', 12))
                ax2.set_ylabel(format_arabic('الرصيد (د.ج)'), fontproperties=FontManager.get_matplotlib_font_prop('Regular', 12))
                ax2.fill_between(all_trans['date'], all_trans['balance'], color='#1E88E5', alpha=0.2)

                tick_font_prop = FontManager.get_matplotlib_font_prop('Regular', 10)
                plt.setp(ax2.get_xticklabels(), fontproperties=tick_font_prop, rotation=45, ha='right')
                plt.setp(ax2.get_yticklabels(), fontproperties=tick_font_prop)

                self.balance_chart_canvas = FigureCanvasTkAgg(fig2, master=self.charts_frame)
                self.balance_chart_canvas.draw()
                self.balance_chart_canvas.get_tk_widget().grid(row=0, column=0, sticky="nsew", padx=(20, 10), pady=10)

    def create_styled_figure(self):
        plt.style.use('seaborn-v0_8-darkgrid')
        fig, ax = plt.subplots(figsize=(6, 4))
        fig.tight_layout(pad=3.0)

        is_dark = ctk.get_appearance_mode() == "Dark"
        bg_color_name = ctk.ThemeManager.theme["CTkFrame"]["fg_color"][1 if is_dark else 0]
        text_color_name = ctk.ThemeManager.theme["CTkLabel"]["text_color"][1 if is_dark else 0]

        def convert_color(color_name):
            try: return f'#{self.winfo_rgb(color_name)[0]//256:02x}{self.winfo_rgb(color_name)[1]//256:02x}{self.winfo_rgb(color_name)[2]//256:02x}'
            except: return "#2b2b2b" if is_dark else "#f0f0f0"

        bg_color_hex = convert_color(bg_color_name)
        text_color_hex = convert_color(text_color_name)

        fig.patch.set_facecolor(bg_color_hex)
        ax.set_facecolor(bg_color_hex)
        ax.tick_params(colors=text_color_hex)
        for spine in ax.spines.values(): spine.set_color(text_color_hex)

        return fig, ax

class ReportsFrame(BaseDataFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(self, text="التقارير والتحليلات", font=FontManager.TITLE_FONT).grid(row=0, column=0, padx=20, pady=20, sticky="e")

        self.main_tabview = ctk.CTkTabview(self, anchor="e")
        self.main_tabview.grid(row=1, column=0, padx=20, pady=10, sticky="nsew")
        self.main_tabview._segmented_button.configure(font=FontManager.BUTTON_FONT)

        self.main_tabview.add("تقارير التبرعات")
        self.main_tabview.add("تقارير الأداء")
        
        # Setup UI for each main tab
        self.setup_donations_tab()
        self.setup_performance_tab()

    def on_show(self):
        # When the main reports frame is shown, update the currently visible sub-tab
        self.update_donations_tab()
        self.update_performance_tab()

    def setup_donations_tab(self):
        donations_tab = self.main_tabview.tab("تقارير التبرعات")
        donations_tab.grid_rowconfigure(0, weight=1)
        donations_tab.grid_columnconfigure(0, weight=1)

        # Sub-tabs for donations
        sub_tabview = ctk.CTkTabview(donations_tab, anchor="e")
        sub_tabview.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        sub_tabview._segmented_button.configure(font=FontManager.APP_FONT)
        
        sub_tabview.add("ملخص التبرعات")
        sub_tabview.add("التبرعات المالية")
        sub_tabview.add("التبرعات العينية")

        # --- Summary Tab ---
        summary_frame = sub_tabview.tab("ملخص التبرعات")
        summary_frame.grid_rowconfigure(0, weight=1)
        summary_frame.grid_columnconfigure(0, weight=1)
        self.summary_tree = self.create_treeview(summary_frame, ("amount", "category", "payer", "date"), ("المبلغ/القيمة", "الفئة", "المتبرع", "التاريخ"))

        # --- Financial Tab ---
        financial_frame = sub_tabview.tab("التبرعات المالية")
        financial_frame.grid_rowconfigure(0, weight=1)
        financial_frame.grid_columnconfigure(0, weight=1)
        self.financial_tree = self.create_treeview(financial_frame, ("amount", "category", "payer", "date"), ("المبلغ", "الفئة", "المتبرع", "التاريخ"))

        # --- In-Kind Tab ---
        inkind_frame = sub_tabview.tab("التبرعات العينية")
        inkind_frame.grid_rowconfigure(0, weight=1)
        inkind_frame.grid_columnconfigure(0, weight=1)
        self.inkind_tree = self.create_treeview(inkind_frame, ("amount", "description", "payer", "date"), ("القيمة التقديرية", "الوصف", "المتبرع", "التاريخ"))

    def setup_performance_tab(self):
        performance_tab = self.main_tabview.tab("تقارير الأداء")
        performance_tab.grid_rowconfigure(0, weight=1)
        performance_tab.grid_columnconfigure(0, weight=1)
        
        # Sub-tabs for performance
        sub_tabview = ctk.CTkTabview(performance_tab, anchor="e")
        sub_tabview.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        sub_tabview._segmented_button.configure(font=FontManager.APP_FONT)

        sub_tabview.add("الأداء الشهري والسنوي")
        sub_tabview.add("التنبؤ المالي")

        # --- Monthly/Annual Performance Tab ---
        trends_frame = sub_tabview.tab("الأداء الشهري والسنوي")
        trends_frame.grid_columnconfigure((0, 1), weight=1)
        trends_frame.grid_rowconfigure(0, weight=1)
        self.monthly_chart_frame = ctk.CTkFrame(trends_frame)
        self.monthly_chart_frame.grid(row=0, column=1, sticky="nsew", padx=(5,0), pady=5)
        self.annual_chart_frame = ctk.CTkFrame(trends_frame)
        self.annual_chart_frame.grid(row=0, column=0, sticky="nsew", padx=(0,5), pady=5)
        self.monthly_canvas, self.annual_canvas = None, None

        # --- Forecast Tab ---
        forecast_frame = sub_tabview.tab("التنبؤ المالي")
        forecast_frame.grid_rowconfigure(0, weight=1)
        forecast_frame.grid_columnconfigure(0, weight=1)
        self.forecast_chart_frame = ctk.CTkFrame(forecast_frame)
        self.forecast_chart_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        self.forecast_canvas = None

    def update_donations_tab(self):
        incomes = self.db.get_transactions('incomes')
        
        # Summary Tree
        self.populate_tree(self.summary_tree, incomes)
        
        # Financial Tree
        financial_donations = [row for row in incomes if row['category'] != 'تبرعات عينية']
        self.populate_tree(self.financial_tree, financial_donations)
        
        # In-kind Tree
        inkind_donations = [row for row in incomes if row['category'] == 'تبرعات عينية']
        self.populate_tree(self.inkind_tree, inkind_donations)

    def update_performance_tab(self):
        incomes = self.db.get_transactions('incomes')
        expenses = self.db.get_transactions('expenses')
        
        df_incomes = pd.DataFrame([dict(row) for row in incomes])
        df_expenses = pd.DataFrame([dict(row) for row in expenses])
        
        # Update Monthly and Annual charts
        self.update_trend_charts(df_incomes, df_expenses)
        
        # Update Forecast chart
        self.update_forecast_chart(df_incomes, df_expenses)

    def update_trend_charts(self, df_incomes, df_expenses):
        if self.monthly_canvas: self.monthly_canvas.get_tk_widget().destroy()
        if self.annual_canvas: self.annual_canvas.get_tk_widget().destroy()

        # Monthly Chart
        if not df_incomes.empty or not df_expenses.empty:
            fig_monthly, ax_monthly = self.create_styled_figure()
            if not df_incomes.empty:
                df_incomes['date'] = pd.to_datetime(df_incomes['date'])
                income_monthly = df_incomes.set_index('date').resample('ME')['amount'].sum()
            else: income_monthly = pd.Series()
            if not df_expenses.empty:
                df_expenses['date'] = pd.to_datetime(df_expenses['date'])
                expense_monthly = df_expenses.set_index('date').resample('ME')['amount'].sum()
            else: expense_monthly = pd.Series()
            
            monthly_summary = pd.DataFrame({'المداخيل': income_monthly, 'المصاريف': expense_monthly}).fillna(0)
            monthly_summary.index = monthly_summary.index.strftime('%Y-%m')
            monthly_summary.plot(kind='bar', ax=ax_monthly, color=['#009688', '#E53935'])
            self.finalize_chart(fig_monthly, ax_monthly, 'الأداء المالي الشهري', 'الشهر', 'المبلغ (د.ج)')
            self.monthly_canvas = FigureCanvasTkAgg(fig_monthly, master=self.monthly_chart_frame)
            self.monthly_canvas.draw()
            self.monthly_canvas.get_tk_widget().pack(fill="both", expand=True)

        # Annual Chart
        if not df_incomes.empty or not df_expenses.empty:
            fig_annual, ax_annual = self.create_styled_figure()
            if not df_incomes.empty:
                income_annual = df_incomes.set_index('date').resample('YE')['amount'].sum()
            else: income_annual = pd.Series()
            if not df_expenses.empty:
                expense_annual = df_expenses.set_index('date').resample('YE')['amount'].sum()
            else: expense_annual = pd.Series()

            annual_summary = pd.DataFrame({'المداخيل': income_annual, 'المصاريف': expense_annual}).fillna(0)
            annual_summary.index = annual_summary.index.strftime('%Y')
            annual_summary.plot(kind='bar', ax=ax_annual, color=['#007ACC', '#C70039'])
            self.finalize_chart(fig_annual, ax_annual, 'الأداء المالي السنوي', 'السنة', 'المبلغ (د.ج)')
            self.annual_canvas = FigureCanvasTkAgg(fig_annual, master=self.annual_chart_frame)
            self.annual_canvas.draw()
            self.annual_canvas.get_tk_widget().pack(fill="both", expand=True)

    def update_forecast_chart(self, df_incomes, df_expenses):
        if self.forecast_canvas: self.forecast_canvas.get_tk_widget().destroy()

        fig, ax = self.create_styled_figure()

        if df_incomes.empty and df_expenses.empty:
            ax.text(0.5, 0.5, format_arabic('لا توجد بيانات كافية للتنبؤ'), horizontalalignment='center', verticalalignment='center', transform=ax.transAxes)
        else:
            df_incomes['amount_signed'] = df_incomes['amount']
            df_expenses['amount_signed'] = -df_expenses['amount']
            all_trans = pd.concat([df_incomes, df_expenses]).sort_values('date')
            all_trans['date'] = pd.to_datetime(all_trans['date'])
            all_trans['balance'] = all_trans['amount_signed'].cumsum()

            # Plot historical data
            all_trans.plot(x='date', y='balance', ax=ax, label=format_arabic('الرصيد الفعلي'))

            # Forecasting
            if len(all_trans) > 2:
                last_90_days = all_trans[all_trans['date'] >= (all_trans['date'].max() - timedelta(days=90))]
                if len(last_90_days) > 2:
                    last_90_days['days_since_start'] = (last_90_days['date'] - last_90_days['date'].min()).dt.days
                    x = last_90_days['days_since_start']
                    y = last_90_days['balance']
                    
                    # Simple linear regression
                    coeffs = np.polyfit(x, y, 1)
                    slope = coeffs[0]
                    intercept = y.iloc[-1] - slope * x.iloc[-1] # Adjust intercept to start from the last point
                    
                    last_date = last_90_days['date'].max()
                    last_balance = last_90_days['balance'].iloc[-1]
                    
                    forecast_dates = pd.to_datetime([last_date + timedelta(days=i) for i in range(1, 61)])
                    forecast_days_since = (forecast_dates - last_date).days
                    
                    forecast_balance = slope * forecast_days_since + last_balance
                    
                    forecast_df = pd.DataFrame({'date': forecast_dates, 'forecast': forecast_balance})
                    forecast_df.plot(x='date', y='forecast', ax=ax, color='red', linestyle='--', label=format_arabic('الرصيد المتوقع'))

            self.finalize_chart(fig, ax, 'التنبؤ بالرصيد المالي (60 يومًا)', 'التاريخ', 'الرصيد (د.ج)')

        self.forecast_canvas = FigureCanvasTkAgg(fig, master=self.forecast_chart_frame)
        self.forecast_canvas.draw()
        self.forecast_canvas.get_tk_widget().pack(fill="both", expand=True)

    def create_treeview(self, parent, columns, headings):
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)
        tree_frame = ctk.CTkFrame(parent)
        tree_frame.grid(row=0, column=0, sticky="nsew")
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)

        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", selectmode="browse")
        for col, head in zip(columns, headings):
            tree.heading(col, text=head, anchor='center')
            tree.column(col, anchor='e', width=150)
        
        tree.grid(row=0, column=0, sticky="nsew")
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.grid(row=0, column=1, sticky="ns")
        return tree

    def populate_tree(self, tree, data):
        tree.delete(*tree.get_children())
        for row_data in data:
            row = dict(row_data)
            values = []
            for col in tree['columns']:
                if col == 'amount':
                    values.append(f"{row.get(col, 0):,.2f}\u200e د.ج")
                else:
                    values.append(row.get(col, ''))
            tree.insert("", "end", values=values)

    def create_styled_figure(self):
        plt.style.use('seaborn-v0_8-darkgrid')
        fig, ax = plt.subplots(figsize=(6, 4))
        fig.tight_layout(pad=4.0)
        is_dark = ctk.get_appearance_mode() == "Dark"
        
        def convert_color(color_name):
            try:
                rgb = self.winfo_rgb(color_name)
                return f'#{rgb[0]//256:02x}{rgb[1]//256:02x}{rgb[2]//256:02x}'
            except:
                return "#2b2b2b" if is_dark else "#f0f0f0"

        bg_color_name = ctk.ThemeManager.theme["CTkFrame"]["fg_color"][1 if is_dark else 0]
        text_color_name = ctk.ThemeManager.theme["CTkLabel"]["text_color"][1 if is_dark else 0]
        
        bg_color_hex = convert_color(bg_color_name)
        text_color_hex = convert_color(text_color_name)

        fig.patch.set_facecolor(bg_color_hex)
        ax.set_facecolor(bg_color_hex)
        ax.tick_params(colors=text_color_hex, which='both')
        ax.xaxis.label.set_color(text_color_hex)
        ax.yaxis.label.set_color(text_color_hex)
        ax.title.set_color(text_color_hex)
        for spine in ax.spines.values():
            spine.set_edgecolor(text_color_hex)
        return fig, ax

    def finalize_chart(self, fig, ax, title, xlabel, ylabel):
        title_prop = FontManager.get_matplotlib_font_prop('Bold', 14)
        label_prop = FontManager.get_matplotlib_font_prop('Regular', 12)
        tick_prop = FontManager.get_matplotlib_font_prop('Regular', 10)
        
        ax.set_title(format_arabic(title), fontproperties=title_prop)
        ax.set_xlabel(format_arabic(xlabel), fontproperties=label_prop)
        ax.set_ylabel(format_arabic(ylabel), fontproperties=label_prop)
        
        if ax.get_legend():
            legend_prop = FontManager.get_matplotlib_font_prop('Regular', 10)
            for text in ax.get_legend().get_texts():
                text.set_fontproperties(legend_prop)
                text.set_text(format_arabic(text.get_text()))

        plt.setp(ax.get_xticklabels(), fontproperties=tick_prop, rotation=30, ha='right')
        plt.setp(ax.get_yticklabels(), fontproperties=tick_prop)
        fig.tight_layout()

class DataManagementFrame(ctk.CTkFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, fg_color="transparent")
        self.controller = controller
        self.grid_columnconfigure(0, weight=1); self.grid_rowconfigure(1, weight=1)
        ctk.CTkLabel(self, text="إدارة البيانات", font=FontManager.TITLE_FONT).grid(row=0, column=0, padx=20, pady=20, sticky="e")

        tabview = ctk.CTkTabview(self, anchor="e"); tabview.grid(row=1, column=0, padx=20, pady=10, sticky="nsew")
        tab_categories_text = "إدارة الفئات"
        tab_backup_text = "النسخ الاحتياطي والمسح"
        tab_settings_text = "إعدادات عامة"

        tabview._segmented_button.configure(font=FontManager.BUTTON_FONT)

        tabview.add(tab_categories_text)
        tabview.add(tab_backup_text)
        tabview.add(tab_settings_text)

        categories_frame = tabview.tab(tab_categories_text); categories_frame.grid_columnconfigure((0, 1), weight=1)
        self.income_cat_manager = CategoryManagementWidget(categories_frame, controller, "income_categories", "فئات المداخيل")
        self.income_cat_manager.grid(row=0, column=1, padx=(10, 20), pady=20, sticky="nsew")
        self.expense_cat_manager = CategoryManagementWidget(categories_frame, controller, "expense_categories", "فئات المصاريف")
        self.expense_cat_manager.grid(row=0, column=0, padx=(20, 10), pady=20, sticky="nsew")

        backup_frame = tabview.tab(tab_backup_text)
        ctk.CTkButton(backup_frame, text="إنشاء نسخة احتياطية كاملة", image=self.controller.icons.get("export"), compound="right", command=self.controller.export_full_backup, font=FontManager.BUTTON_FONT).pack(fill="x", padx=20, pady=(20, 10))
        ctk.CTkButton(backup_frame, text="استعادة بيانات من نسخة احتياطية", image=self.controller.icons.get("import"), compound="right", command=self.controller.restore_from_excel, font=FontManager.BUTTON_FONT).pack(fill="x", padx=20, pady=10)
        ctk.CTkButton(backup_frame, text="مسح جميع البيانات المالية", fg_color="#D2042D", hover_color="#990000", font=FontManager.BUTTON_FONT, command=self.clear_all_data).pack(fill="x", padx=20, pady=(50, 10))
        ctk.CTkLabel(backup_frame, text="تحذير: مسح البيانات سيقوم بحذف جميع المداخيل والمصاريف بشكل نهائي.", wraplength=400, justify="right", font=FontManager.APP_FONT).pack(fill="x", padx=20)

        settings_frame = tabview.tab(tab_settings_text)
        settings_scroll_frame = ctk.CTkScrollableFrame(settings_frame, label_text="إعدادات الوصل والبرنامج", label_font=FontManager.H2_FONT)
        settings_scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        settings_scroll_frame.grid_columnconfigure(1, weight=1)
        self.setting_entries = {}
        settings_fields = {
            'association_name': "اسم الجمعية:",
            'address': "العنوان:",
            'phone': "رقم الهاتف:"
        }
        for i, (key, label_text) in enumerate(settings_fields.items()):
            ctk.CTkLabel(settings_scroll_frame, text=label_text, font=FontManager.APP_FONT).grid(row=i, column=0, padx=20, pady=(15,0), sticky="e")
            entry = ctk.CTkEntry(settings_scroll_frame, justify='right', font=FontManager.INPUT_FONT)
            entry.grid(row=i, column=1, padx=20, pady=(15,0), sticky="ew")
            self.setting_entries[key] = entry

        save_button = ctk.CTkButton(settings_scroll_frame, text="حفظ الإعدادات", command=self.save_settings, font=FontManager.BUTTON_FONT)
        save_button.grid(row=len(settings_fields), column=1, padx=20, pady=20, sticky="w")
        tabview.set(tab_categories_text)

    def on_show(self):
        self.income_cat_manager.populate_table()
        self.expense_cat_manager.populate_table()
        self.load_settings()

    def load_settings(self):
        all_settings = self.controller.db.get_all_settings()
        for key, entry in self.setting_entries.items():
            entry.delete(0, 'end')
            entry.insert(0, all_settings.get(key, ''))

    def save_settings(self):
        try:
            for key, entry in self.setting_entries.items():
                self.controller.db.update_setting(key, entry.get())
            messagebox.showinfo("نجاح", "تم حفظ الإعدادات بنجاح.", parent=self)
            self.controller.db.log_action(self.controller.current_user, "تحديث الإعدادات العامة")
        except Exception as e:
            messagebox.showerror("خطأ", f"فشل حفظ الإعدادات:\n{e}", parent=self)

    def clear_all_data(self):
        if messagebox.askyesno("تأكيد المسح", "هل أنت متأكد من مسح جميع السجلات المالية؟\nلا يمكن التراجع عن هذا الإجراء.", icon='warning', parent=self):
            self.controller.db.clear_all_transactions()
            try:
                if os.path.exists(self.controller.attachments_dir): shutil.rmtree(self.controller.attachments_dir)
                os.makedirs(self.controller.attachments_dir)
            except OSError as e:
                messagebox.showerror("خطأ", f"فشل حذف المرفقات:\n{e}", parent=self)

            self.controller.db.log_action(self.controller.current_user, "مسح جميع البيانات")
            messagebox.showinfo("نجاح", "تم مسح جميع البيانات والمرفقات.", parent=self)
            self.controller.refresh_all_data()

class CategoryManagementWidget(ctk.CTkFrame):
    def __init__(self, parent, controller, table_name, title):
        super().__init__(parent); self.controller, self.db, self.table_name = controller, controller.db, table_name
        self.grid_columnconfigure(0, weight=1); self.grid_rowconfigure(1, weight=1)
        top_frame = ctk.CTkFrame(self, fg_color="transparent"); top_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=(10,0))
        ctk.CTkLabel(top_frame, text=title, font=FontManager.H2_FONT).pack(side="right")
        button_frame = ctk.CTkFrame(top_frame, fg_color="transparent"); button_frame.pack(side="left")
        ctk.CTkButton(button_frame, text="", image=self.controller.icons.get("add"), width=30, command=self.add_item).pack(side="left", padx=2)
        ctk.CTkButton(button_frame, text="", image=self.controller.icons.get("edit"), width=30, command=self.edit_item).pack(side="left", padx=2)
        ctk.CTkButton(button_frame, text="", image=self.controller.icons.get("delete"), width=30, command=self.delete_item).pack(side="left", padx=2)
        tree_frame = ctk.CTkFrame(self); tree_frame.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        tree_frame.grid_rowconfigure(0, weight=1); tree_frame.grid_columnconfigure(0, weight=1)
        self.tree = ttk.Treeview(tree_frame, columns=("id", "name"), show="headings", selectmode="browse")
        self.tree.heading("id", text="المعرف"); self.tree.heading("name", text="الاسم"); self.tree.column("id", width=60, anchor="center"); self.tree.column("name", anchor="e")
        self.tree.grid(row=0, column=0, sticky="nsew")

    def populate_table(self):
        self.tree.delete(*self.tree.get_children())
        for cat in self.db.get_categories(self.table_name): self.tree.insert("", "end", values=(cat['id'], cat['name']))

    def add_item(self):
        new_name = CategoryDialog(self, "إضافة فئة جديدة").get_data()
        if new_name:
            if self.db.add_category(self.table_name, new_name):
                self.populate_table(); self.controller.db.log_action(self.controller.current_user, "إضافة فئة", f"الاسم: {new_name}")
            else: messagebox.showerror("خطأ", "هذه الفئة موجودة بالفعل.", parent=self)

    def edit_item(self):
        selected = self.tree.focus()
        if not selected: messagebox.showwarning("تنبيه", "يرجى تحديد فئة لتعديلها.", parent=self); return
        item = self.tree.item(selected)['values']; cat_id, old_name = item[0], item[1]
        new_name = CategoryDialog(self, "تعديل الفئة", initial_value=old_name).get_data()
        if new_name and new_name != old_name:
            if self.db.update_category(self.table_name, cat_id, new_name):
                self.populate_table(); self.controller.db.log_action(self.controller.current_user, "تعديل فئة", f"من '{old_name}' إلى '{new_name}'")
            else: messagebox.showerror("خطأ", "هذه الفئة موجودة بالفعل.", parent=self)

    def delete_item(self):
        selected = self.tree.focus()
        if not selected: messagebox.showwarning("تنبيه", "يرجى تحديد فئة لحذفها.", parent=self); return
        item = self.tree.item(selected)['values']; cat_id, name = item[0], item[1]
        if messagebox.askyesno("تأكيد الحذف", f"هل أنت متأكد من حذف الفئة '{name}'؟\nسيؤثر هذا على التقارير القديمة.", icon='warning', parent=self):
            self.db.delete_category(self.table_name, cat_id); self.populate_table(); self.controller.db.log_action(self.controller.current_user, "حذف فئة", f"الاسم: {name}")

class CategoryDialog(ctk.CTkToplevel):
    def __init__(self, parent, title, initial_value=""):
        super().__init__(parent); self.title(title); self.geometry("400x150"); self.resizable(False, False); self.transient(parent)
        self.data = None
        ctk.CTkLabel(self, text="اسم الفئة:", font=FontManager.APP_FONT).pack(pady=(10,5))
        self.entry = ctk.CTkEntry(self, width=300, justify='right', font=FontManager.INPUT_FONT); self.entry.pack(pady=5, padx=20); self.entry.insert(0, initial_value)
        button_frame = ctk.CTkFrame(self, fg_color="transparent"); button_frame.pack(pady=10)
        ctk.CTkButton(button_frame, text="حفظ", command=self.save, font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        ctk.CTkButton(button_frame, text="إلغاء", command=self.cancel, fg_color="gray", font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        self.grab_set(); self.entry.focus(); self.wait_window()
    def save(self):
        value = self.entry.get().strip()
        if not value: messagebox.showerror("خطأ", "اسم الفئة لا يمكن أن يكون فارغاً.", parent=self); return
        self.data = value; self.destroy()
    def cancel(self): self.data = None; self.destroy()
    def get_data(self): return self.data

class AboutWindow(ctk.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("حول البرنامج")
        self.geometry("450x450")
        self.resizable(False, False)
        self.transient(parent)
        self.logo_image = None

        try:
            try: base_path = sys._MEIPASS
            except AttributeError: base_path = os.path.abspath(".")
            logo_path = os.path.join(base_path, "Logo FM.png")

            logo_image_pil = Image.open(logo_path)
            logo_image_pil.thumbnail((120, 120))
            self.logo_image = ctk.CTkImage(light_image=logo_image_pil, dark_image=logo_image_pil, size=(logo_image_pil.width, logo_image_pil.height))
            logo_label = ctk.CTkLabel(self, image=self.logo_image, text="")
            logo_label.pack(pady=(20, 10))
        except Exception as e:
            print(f"Error loading logo: {e}")
            logo_label = ctk.CTkLabel(self, text="لم يتم العثور على الشعار", font=FontManager.APP_FONT)
            logo_label.pack(pady=(20, 10))


        ctk.CTkLabel(self, text="برنامج إدارة الجمعية", font=FontManager.TITLE_FONT).pack(pady=(0, 10))
        ctk.CTkLabel(self, text=f"الإصدار: {APP_VERSION}", font=FontManager.APP_FONT).pack()
        ctk.CTkLabel(self, text="تم تطوير هذا البرنامج لتسهيل الإدارة المالية لجمعيات المساجد.", font=FontManager.APP_FONT, wraplength=400, justify="center").pack(pady=20)
        ctk.CTkLabel(self, text="fakloumohammed@gmail.com :للتواصل والدعم ", font=FontManager.SMALL_FONT).pack()
        ctk.CTkButton(self, text="موافق", command=self.destroy, font=FontManager.BUTTON_FONT).pack(pady=(20,10))
        self.grab_set()
        self.focus()

class UserManagementFrame(BaseDataFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        top_frame = ctk.CTkFrame(self); top_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        ctk.CTkLabel(top_frame, text="إدارة المستخدمين", font=FontManager.TITLE_FONT).pack(side="right", padx=10, pady=5)
        ctk.CTkButton(top_frame, text="إضافة", image=self.controller.icons.get("add"), compound="right", command=self.add_user, font=FontManager.BUTTON_FONT).pack(side="left", padx=5, pady=10)
        ctk.CTkButton(top_frame, text="تعديل", image=self.controller.icons.get("edit"), compound="right", command=self.edit_user, font=FontManager.BUTTON_FONT).pack(side="left", padx=5, pady=10)
        ctk.CTkButton(top_frame, text="حذف", image=self.controller.icons.get("delete"), compound="right", command=self.delete_user, fg_color="#D2042D", hover_color="#990000", font=FontManager.BUTTON_FONT).pack(side="left", padx=5, pady=10)
        self.setup_treeview(("id", "username", "role"), ("المعرف", "اسم المستخدم", "الصلاحية"))
    def on_show(self): self.populate_table()
    def populate_table(self):
        for i in self.tree.get_children(): self.tree.delete(i)
        for user in self.db.get_all_users(): self.tree.insert("", "end", values=(user['id'], user['username'], user['role']))
    def add_user(self):
        data = UserDialog(self, title="إضافة مستخدم جديد").get_data()
        if data and self.db.add_user(data['username'], data['password'], data['role']):
            self.db.log_action(self.controller.current_user, "إضافة مستخدم", f"اسم: {data['username']}"); self.populate_table()
        else: messagebox.showerror("خطأ", "اسم المستخدم موجود بالفعل.", parent=self)
    def edit_user(self):
        selected = self.tree.focus()
        if not selected: messagebox.showwarning("تنبيه", "يرجى تحديد مستخدم.", parent=self); return
        vals = self.tree.item(selected)['values']; user_id, username, role = vals[0], vals[1], vals[2]
        data = UserDialog(self, title="تعديل مستخدم", initial_data={'username': username, 'role': role}).get_data()
        if data: self.db.update_user(user_id, data['username'], data['password'], data['role']); self.db.log_action(self.controller.current_user, "تعديل مستخدم", f"معرف: {user_id}"); self.populate_table()
    def delete_user(self):
        selected = self.tree.focus()
        if not selected: messagebox.showwarning("تنبيه", "يرجى تحديد مستخدم.", parent=self); return
        vals = self.tree.item(selected)['values']; user_id, username = vals[0], vals[1]
        if user_id == 1: messagebox.showerror("خطأ", "لا يمكن حذف المدير الافتراضي.", parent=self); return
        if username == self.controller.current_user: messagebox.showerror("خطأ", "لا يمكن للمستخدم حذف نفسه.", parent=self); return
        if messagebox.askyesno("تأكيد الحذف", f"هل أنت متأكد من حذف المستخدم '{username}'؟", icon='warning', parent=self):
            self.db.delete_user(user_id); self.db.log_action(self.controller.current_user, "حذف مستخدم", f"معرف: {user_id}"); self.populate_table()

class AuditLogFrame(BaseDataFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        ctk.CTkLabel(self, text="سجل التدقيق", font=FontManager.TITLE_FONT).grid(row=0, column=0, padx=20, pady=20, sticky="e")
        columns = ("timestamp", "username", "action", "details"); headings = ("الوقت والتاريخ", "المستخدم", "الإجراء", "التفاصيل")
        self.setup_treeview(columns, headings); self.tree.column("timestamp", width=160, anchor="center"); self.tree.column("details", width=400); self.tree.column("action", width=150); self.tree.column("username", width=120)
    def on_show(self):
        for i in self.tree.get_children(): self.tree.delete(i)
        for log in self.db.get_audit_logs(): self.tree.insert("", "end", values=(log['timestamp'], log['username'], log['action'], log['details']))

class UserDialog(ctk.CTkToplevel):
    def __init__(self, parent, title, initial_data=None):
        super().__init__(parent); self.title(title); self.geometry("450x350"); self.resizable(False, False); self.transient(parent); self.protocol("WM_DELETE_WINDOW", self.cancel)
        self.data, self.entries = None, {}
        main_frame = ctk.CTkFrame(self, fg_color="transparent"); main_frame.pack(expand=True, padx=20, pady=20); main_frame.grid_columnconfigure(1, weight=1)
        fields = {"username": "اسم المستخدم", "password": "كلمة المرور", "role": "الصلاحية"}
        for i, (key, label_text) in enumerate(fields.items()):
            ctk.CTkLabel(main_frame, text=label_text, font=FontManager.APP_FONT).grid(row=i, column=0, padx=10, pady=10, sticky="e")
            if key == "role": entry = ctk.CTkComboBox(main_frame, values=["مدير", "أمين مال"], state="readonly", justify="right", font=FontManager.INPUT_FONT, dropdown_font=FontManager.INPUT_FONT)
            else: entry = ctk.CTkEntry(main_frame, width=250, show=("*" if key == "password" else ""), justify="right", font=FontManager.INPUT_FONT)
            entry.grid(row=i, column=1, padx=10, pady=10, sticky="ew"); self.entries[key] = entry
        if initial_data: self.entries['username'].insert(0, initial_data['username']); self.entries['role'].set(initial_data['role']); self.entries['password'].configure(placeholder_text="اتركها فارغة لعدم التغيير")
        else: self.entries['role'].set("أمين مال")
        ctk.CTkButton(main_frame, text="حفظ", command=self.save, font=FontManager.BUTTON_FONT).grid(row=len(fields), column=0, columnspan=2, pady=20)
        self.grab_set(); self.wait_window()
    def save(self):
        username, password, role = self.entries['username'].get().strip(), self.entries['password'].get().strip(), self.entries['role'].get()
        if not username or not role: messagebox.showerror("خطأ", "يجب إدخال اسم المستخدم والصلاحية.", parent=self); return
        if "إضافة" in self.title() and not password: messagebox.showerror("خطأ", "كلمة المرور مطلوبة عند إضافة مستخدم جديد.", parent=self); return
        self.data = {'username': username, 'password': password, 'role': role}; self.destroy()
    def cancel(self): self.data = None; self.destroy()
    def get_data(self): return self.data

class ChangePasswordDialog(ctk.CTkToplevel):
    def __init__(self, parent, user_id):
        super().__init__(parent); self.controller, self.user_id = parent, user_id
        self.title("تغيير كلمة المرور الإلزامي"); self.geometry("450x380"); self.resizable(False, False); self.transient(parent); self.protocol("WM_DELETE_WINDOW", self.force_change)
        frame = ctk.CTkFrame(self, fg_color="transparent"); frame.pack(expand=True, padx=20, pady=20)
        ctk.CTkLabel(frame, text="مرحباً بك أيها المدير", font=FontManager.TITLE_FONT).pack(pady=(0,10))
        ctk.CTkLabel(frame, text="لأسباب أمنية، يجب عليك تغيير كلمة المرور الافتراضية.", wraplength=350, justify='center', font=FontManager.APP_FONT).pack(pady=(0, 20))
        ctk.CTkLabel(frame, text="كلمة المرور الجديدة", font=FontManager.APP_FONT).pack()
        self.new_pass_entry = ctk.CTkEntry(frame, width=250, show="*", justify='center', font=FontManager.INPUT_FONT); self.new_pass_entry.pack(pady=5)
        ctk.CTkLabel(frame, text="تأكيد كلمة المرور الجديدة", font=FontManager.APP_FONT).pack()
        self.confirm_pass_entry = ctk.CTkEntry(frame, width=250, show="*", justify='center', font=FontManager.INPUT_FONT); self.confirm_pass_entry.pack(pady=5)
        ctk.CTkButton(frame, text="حفظ وتغيير كلمة المرور", command=self.save_password, font=FontManager.BUTTON_FONT, width=250).pack(pady=20)
        self.grab_set(); self.wait_window()
    def force_change(self): messagebox.showwarning("إجراء إلزامي", "يجب عليك تغيير كلمة المرور للمتابعة.", parent=self); self.wait_window()
    def save_password(self):
        new_pass, confirm_pass = self.new_pass_entry.get(), self.confirm_pass_entry.get()
        if not new_pass or not confirm_pass: messagebox.showerror("خطأ", "الرجاء ملء كلا الحقلين.", parent=self); return
        if new_pass != confirm_pass: messagebox.showerror("خطأ", "كلمتا المرور غير متطابقتين.", parent=self); return
        if new_pass == "admin": messagebox.showerror("خطأ", "الرجاء اختيار كلمة مرور مختلفة عن الكلمة الافتراضية.", parent=self); return
        user_data = self.controller.db.get_user("admin")
        if user_data:
            self.controller.db.update_user(self.user_id, "admin", new_pass, "مدير", must_change_password=0)
            self.controller.db.log_action("admin", "تغيير كلمة المرور الإلزامي")
            messagebox.showinfo("نجاح", "تم تغيير كلمة المرور بنجاح.", parent=self)
            self.destroy()
        else:
             messagebox.showerror("خطأ", "لم يتم العثور على المستخدم admin.", parent=self)
             self.destroy()

class WordReceiptGenerator:
    def __init__(self, receipt_data, settings):
        self.receipt_data = receipt_data
        self.settings = settings
        self.document = docx.Document()

        section = self.document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5) # Adjusted for better layout
        section.left_margin = section.right_margin = Cm(1.5)
        self.document.styles['Normal'].font.rtl = True
        self.document.styles['Normal'].font.name = 'Amiri'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Amiri')


    def _set_font(self, run, font_name='Amiri', size=None, bold=None, italic=None):
        font = run.font
        font.name = font_name
        r = run._r
        r.rPr.rFonts.set(qn('w:cs'), font_name)
        r.rPr.rFonts.set(qn('w:rtl'), font_name)
        if size: font.size = Pt(size)
        if bold is not None: font.bold = bold
        if italic is not None: font.italic = italic

    def _get_or_create_tblPr(self, tbl_element):
        tblPr = tbl_element.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_element.insert(0, tblPr)
        return tblPr

    def _remove_table_borders(self, table):
        tblPr = self._get_or_create_tblPr(table._element)
        tblBorders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            ele = OxmlElement(f'w:{b}')
            ele.set(qn('w:val'), 'none')
            tblBorders.append(ele)
        tblPr.append(tblBorders)

    def create_document(self):
        # This creates a single receipt. To create two, we wrap this logic.
        self.create_receipt_content(is_top_part=True)

        # --- Separator Line ---
        self.document.add_paragraph("------------------------------------------------------------------------------------------------------------------")

        self.create_receipt_content(is_top_part=False)


    def create_receipt_content(self, is_top_part=True):
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("بسم الله الرحمن الرحيم")
        self._set_font(run, size=16, bold=True)

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("وصل استلام تبرع")
        self._set_font(run, size=22, bold=True)

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.settings.get("association_name", "اسم الجمعية"))
        self._set_font(run, size=12, bold=True)

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address = self.settings.get("address", "(العنوان)")
        phone = self.settings.get("phone", "(الهاتف)")
        run = p.add_run(f"العنوان: {address} - الهاتف: {phone}")
        self._set_font(run, size=10)

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("__________________________________________")
        self._set_font(run, size=10)

        table = self.document.add_table(rows=1, cols=2)
        table.columns[0].width = table.columns[1].width = Cm(8)
        self._remove_table_borders(table)

        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"رقم الوصل: {self.receipt_data.get('id', 'N/A')}")
        self._set_font(run, size=12, bold=True)

        p = table.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"تاريخ التحرير: {self.receipt_data.get('date', datetime.now().strftime('%d/%m/%Y'))}")
        self._set_font(run, size=12, bold=True)

        self.document.add_paragraph()

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"استلمنا من السيد/ة: {self.receipt_data.get('payer', 'N/A')}")
        self._set_font(run, size=12)

        amount = self.receipt_data.get('amount', 0)
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"مبلغا وقدره: {amount:,.2f} د.ج")
        self._set_font(run, size=12)

        amount_text = tafqeet(amount)
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"وهو: ({amount_text} دينار جزائري)")
        self._set_font(run, size=12)

        description = self.receipt_data.get('description', 'تبرع نقدي')
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"وذلك عن: {description}")
        self._set_font(run, size=12)

        self.document.add_paragraph() # Spacer

        table = self.document.add_table(rows=1, cols=2)
        table.columns[0].width = table.columns[1].width = Cm(8)
        self._remove_table_borders(table)

        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ختم الجمعية")
        self._set_font(run, size=14, bold=True)

        p = table.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("توقيع أمين المال")
        self._set_font(run, size=14, bold=True)

        self.document.add_paragraph()

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = "نسخة للجمعية" if is_top_part else "نسخة للمتبرع"
        run = p.add_run(f"({note}) لا يعتبر هذا الوصل صالحاً إلا إذا كان مختوماً وموقعاً.")
        self._set_font(run, size=9, italic=True)


    def save(self, filepath):
        try:
            self.create_document()
            folder = os.path.dirname(filepath)
            if folder and not os.path.exists(folder):
                os.makedirs(folder)
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except PermissionError:
                    base, ext = os.path.splitext(filepath)
                    filepath = f"{base}_جديد{ext}"
            self.document.save(filepath)
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("خطأ", f"حدث خطأ غير متوقع أثناء حفظ الوصل:\n{e}")



class MemberDialog(ctk.CTkToplevel):
    def __init__(self, parent, title, initial_data=None):
        super().__init__(parent)
        self.title(title)
        self.geometry("600x450")
        self.resizable(False, False)
        self.transient(parent)
        self.protocol("WM_DELETE_WINDOW", self.cancel)
        self.data = None
        self.entries = {}

        main_frame = ctk.CTkFrame(self, fg_color="transparent")
        main_frame.pack(expand=True, padx=20, pady=20)
        main_frame.grid_columnconfigure(1, weight=1)

        fields = {
            "full_name": {"label": "الاسم الكامل *", "type": "entry"},
            "join_date": {"label": "تاريخ الانضمام *", "type": "entry"},
            "phone": {"label": "رقم الهاتف", "type": "entry"},
            "address": {"label": "العنوان", "type": "entry"},
            "status": {"label": "الحالة", "type": "combo", "values": ["نشط", "غير نشط", "مجمد"]},
            "notes": {"label": "ملاحظات", "type": "text"}
        }

        for i, (key, props) in enumerate(fields.items()):
            ctk.CTkLabel(main_frame, text=props['label'], font=FontManager.APP_FONT).grid(row=i, column=0, padx=10, pady=10, sticky="e")
            if props['type'] == "entry":
                entry = ctk.CTkEntry(main_frame, width=350, justify="right", font=FontManager.INPUT_FONT)
            elif props['type'] == "combo":
                entry = ctk.CTkComboBox(main_frame, values=props['values'], width=350, justify="right", state="readonly", font=FontManager.INPUT_FONT, dropdown_font=FontManager.INPUT_FONT)
            elif props['type'] == "text":
                entry = ctk.CTkTextbox(main_frame, width=350, height=80, font=FontManager.INPUT_FONT)
            
            entry.grid(row=i, column=1, padx=10, pady=10, sticky="ew")
            self.entries[key] = entry

        if not initial_data:
            self.entries['join_date'].insert(0, datetime.now().strftime("%Y-%m-%d"))
            self.entries['status'].set("نشط")
        
        if initial_data:
            for key, widget in self.entries.items():
                value = initial_data.get(key)
                if value is not None:
                    if isinstance(widget, ctk.CTkTextbox):
                        widget.insert("1.0", value)
                    elif isinstance(widget, ctk.CTkComboBox):
                        widget.set(str(value))
                    else:
                        widget.delete(0, "end")
                        widget.insert(0, str(value))

        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.pack(pady=20)
        ctk.CTkButton(button_frame, text="حفظ", command=self.save, font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        ctk.CTkButton(button_frame, text="إلغاء", command=self.cancel, fg_color="gray", font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        
        self.grab_set()
        self.wait_window()

    def save(self):
        self.data = {}
        for key, widget in self.entries.items():
            if isinstance(widget, ctk.CTkTextbox):
                self.data[key] = widget.get("1.0", "end-1c").strip()
            else:
                self.data[key] = widget.get().strip()

        if not self.data['full_name'] or not self.data['join_date']:
            messagebox.showerror("خطأ", "الاسم الكامل وتاريخ الانضمام حقول إلزامية.", parent=self)
            self.data = None
            return
        self.destroy()

    def cancel(self):
        self.data = None
        self.destroy()

    def get_data(self):
        return self.data

class MemberManagementFrame(BaseDataFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        top_frame = ctk.CTkFrame(self)
        top_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        ctk.CTkLabel(top_frame, text="إدارة أعضاء الجمعية", font=FontManager.TITLE_FONT).pack(side="right", padx=10, pady=5)
        
        button_panel = ctk.CTkFrame(top_frame, fg_color="transparent")
        button_panel.pack(side="left", padx=5, pady=10)
        ctk.CTkButton(button_panel, text="إضافة عضو", image=self.controller.icons.get("add"), compound="right", command=self.add_member, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        ctk.CTkButton(button_panel, text="تعديل", image=self.controller.icons.get("edit"), compound="right", command=self.edit_member, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        ctk.CTkButton(button_panel, text="حذف", image=self.controller.icons.get("delete"), compound="right", command=self.delete_member, fg_color="#D2042D", hover_color="#990000", font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        
        columns = ("status", "phone", "join_date", "full_name", "id")
        headings = ("الحالة", "رقم الهاتف", "تاريخ الانضمام", "الاسم الكامل", "المعرّف")
        self.setup_treeview(columns, headings)
        self.tree.column("id", width=60, anchor='center')
        self.tree.column("full_name", width=250, anchor='e')
        self.tree.column("join_date", width=120, anchor='center')
        self.tree.column("phone", width=150, anchor='center')
        self.tree.column("status", width=100, anchor='center')
        
    def on_show(self):
        self.populate_table()

    def populate_table(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        for member in self.db.get_all_members():
            m = dict(member)
            self.tree.insert("", "end", values=(m.get('status'), m.get('phone'), m.get('join_date'), m.get('full_name'), m.get('id')))
    
    def add_member(self):
        data = MemberDialog(self, title="إضافة عضو جديد").get_data()
        if data:
            self.db.add_member(data)
            self.db.log_action(self.controller.current_user, "إضافة عضو", f"اسم العضو: {data['full_name']}")
            self.populate_table()

    def edit_member(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("تنبيه", "يرجى تحديد عضو لتعديله.", parent=self)
            return
        member_id = self.tree.item(selected_item)['values'][-1]
        member_data = self.db.get_member_by_id(member_id)
        if not member_data: return

        data = MemberDialog(self, title="تعديل بيانات عضو", initial_data=dict(member_data)).get_data()
        if data:
            self.db.update_member(member_id, data)
            self.db.log_action(self.controller.current_user, "تعديل عضو", f"معرّف العضو: {member_id}")
            self.populate_table()

    def delete_member(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("تنبيه", "يرجى تحديد عضو لحذفه.", parent=self)
            return
        
        values = self.tree.item(selected_item)['values']
        member_id = values[-1]
        member_name = values[-2]

        if messagebox.askyesno("تأكيد الحذف", f"هل أنت متأكد من حذف العضو '{member_name}'؟", icon='warning', parent=self):
            self.db.delete_member(member_id)
            self.db.log_action(self.controller.current_user, "حذف عضو", f"معرّف العضو: {member_id}")
            self.populate_table()

class ActivityDialog(ctk.CTkToplevel):
    def __init__(self, parent, title, initial_data=None):
        super().__init__(parent)
        self.title(title)
        self.geometry("600x400")
        self.resizable(False, False)
        self.transient(parent)
        self.protocol("WM_DELETE_WINDOW", self.cancel)
        self.data = None
        self.entries = {}

        main_frame = ctk.CTkFrame(self, fg_color="transparent")
        main_frame.pack(expand=True, padx=20, pady=20)
        main_frame.grid_columnconfigure(1, weight=1)

        fields = {
            "name": {"label": "اسم النشاط *", "type": "entry"},
            "date": {"label": "تاريخ النشاط *", "type": "entry"},
            "location": {"label": "المكان", "type": "entry"},
            "description": {"label": "الوصف", "type": "text"}
        }

        for i, (key, props) in enumerate(fields.items()):
            ctk.CTkLabel(main_frame, text=props['label'], font=FontManager.APP_FONT).grid(row=i, column=0, padx=10, pady=10, sticky="e")
            if props['type'] == "entry":
                entry = ctk.CTkEntry(main_frame, width=350, justify="right", font=FontManager.INPUT_FONT)
            elif props['type'] == "text":
                entry = ctk.CTkTextbox(main_frame, width=350, height=100, font=FontManager.INPUT_FONT)
            
            entry.grid(row=i, column=1, padx=10, pady=10, sticky="ew")
            self.entries[key] = entry

        if not initial_data:
            self.entries['date'].insert(0, datetime.now().strftime("%Y-%m-%d"))
        
        if initial_data:
            for key, widget in self.entries.items():
                value = initial_data.get(key)
                if value is not None:
                    if isinstance(widget, ctk.CTkTextbox):
                        widget.insert("1.0", value)
                    else:
                        widget.delete(0, "end")
                        widget.insert(0, str(value))
        
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.grid(row=len(fields), column=0, columnspan=2, pady=20)
        ctk.CTkButton(button_frame, text="حفظ", command=self.save, font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        ctk.CTkButton(button_frame, text="إلغاء", command=self.cancel, fg_color="gray", font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        
        self.grab_set()
        self.wait_window()

    def save(self):
        self.data = {}
        for key, widget in self.entries.items():
            if isinstance(widget, ctk.CTkTextbox):
                self.data[key] = widget.get("1.0", "end-1c").strip()
            else:
                self.data[key] = widget.get().strip()

        if not self.data['name'] or not self.data['date']:
            messagebox.showerror("خطأ", "اسم النشاط وتاريخه حقول إلزامية.", parent=self)
            self.data = None
            return
        self.destroy()

    def cancel(self):
        self.data = None
        self.destroy()

    def get_data(self):
        return self.data

class AttendanceDialog(ctk.CTkToplevel):
    def __init__(self, parent, controller, activity_id, activity_name):
        super().__init__(parent)
        self.controller = controller
        self.activity_id = activity_id
        self.title(f"تسجيل الحضور - {activity_name}")
        self.geometry("500x600")
        self.transient(parent)

        self.member_vars = {}

        ctk.CTkLabel(self, text=f"اختر الأعضاء الحاضرين في نشاط:", font=FontManager.H2_FONT).pack(pady=(10,0))
        ctk.CTkLabel(self, text=activity_name, font=FontManager.H1_FONT).pack(pady=(0,10))

        scrollable_frame = ctk.CTkScrollableFrame(self)
        scrollable_frame.pack(fill="both", expand=True, padx=20, pady=10)

        all_members = self.controller.db.get_all_members()
        present_member_ids = self.controller.db.get_attendance(self.activity_id)

        for member in all_members:
            member_id = member['id']
            var = ctk.StringVar(value="on" if member_id in present_member_ids else "off")
            cb = ctk.CTkCheckBox(scrollable_frame, text=member['full_name'], variable=var, onvalue="on", offvalue="off", font=FontManager.APP_FONT)
            cb.pack(anchor="e", padx=10, pady=5)
            self.member_vars[member_id] = var

        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.pack(pady=10)
        ctk.CTkButton(button_frame, text="حفظ الحضور", command=self.save_attendance, font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        ctk.CTkButton(button_frame, text="إلغاء", command=self.destroy, fg_color="gray", font=FontManager.BUTTON_FONT).pack(side="left", padx=10)
        
        self.grab_set()

    def save_attendance(self):
        present_ids = [member_id for member_id, var in self.member_vars.items() if var.get() == "on"]
        self.controller.db.update_attendance(self.activity_id, present_ids)
        self.controller.db.log_action(self.controller.current_user, "تحديث الحضور", f"معرف النشاط: {self.activity_id}")
        messagebox.showinfo("نجاح", "تم تحديث قائمة الحضور بنجاح.", parent=self)
        self.destroy()

class ActivityManagementFrame(BaseDataFrame):
    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        top_frame = ctk.CTkFrame(self)
        top_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        ctk.CTkLabel(top_frame, text="إدارة الأنشطة والفعاليات", font=FontManager.TITLE_FONT).pack(side="right", padx=10, pady=5)
        
        button_panel = ctk.CTkFrame(top_frame, fg_color="transparent")
        button_panel.pack(side="left", padx=5, pady=10)
        ctk.CTkButton(button_panel, text="إدارة الحضور", command=self.manage_attendance, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        ctk.CTkButton(button_panel, text="إضافة نشاط", image=self.controller.icons.get("add"), compound="right", command=self.add_activity, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        ctk.CTkButton(button_panel, text="تعديل", image=self.controller.icons.get("edit"), compound="right", command=self.edit_activity, font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        ctk.CTkButton(button_panel, text="حذف", image=self.controller.icons.get("delete"), compound="right", command=self.delete_activity, fg_color="#D2042D", hover_color="#990000", font=FontManager.BUTTON_FONT).pack(side="left", padx=5)
        
        columns = ("description", "location", "date", "name", "id")
        headings = ("الوصف", "المكان", "التاريخ", "اسم النشاط", "المعرّف")
        self.setup_treeview(columns, headings)
        self.tree.column("id", width=60, anchor='center')
        self.tree.column("name", width=250, anchor='e')
        self.tree.column("date", width=120, anchor='center')
        self.tree.column("location", width=200, anchor='e')
        self.tree.column("description", width=350, anchor='e')

    def on_show(self):
        self.populate_table()

    def populate_table(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        for activity in self.db.get_all_activities():
            act = dict(activity)
            self.tree.insert("", "end", values=(act.get('description'), act.get('location'), act.get('date'), act.get('name'), act.get('id')))

    def add_activity(self):
        data = ActivityDialog(self, title="إضافة نشاط جديد").get_data()
        if data:
            self.db.add_activity(data)
            self.db.log_action(self.controller.current_user, "إضافة نشاط", f"اسم النشاط: {data['name']}")
            self.populate_table()

    def edit_activity(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("تنبيه", "يرجى تحديد نشاط لتعديله.", parent=self)
            return
        activity_id = self.tree.item(selected_item)['values'][-1]
        activity_data = self.db.get_activity_by_id(activity_id)
        if not activity_data: return

        data = ActivityDialog(self, title="تعديل بيانات نشاط", initial_data=dict(activity_data)).get_data()
        if data:
            self.db.update_activity(activity_id, data)
            self.db.log_action(self.controller.current_user, "تعديل نشاط", f"معرّف النشاط: {activity_id}")
            self.populate_table()

    def delete_activity(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("تنبيه", "يرجى تحديد نشاط لحذفه.", parent=self)
            return
        
        values = self.tree.item(selected_item)['values']
        activity_id = values[-1]
        activity_name = values[-2]

        if messagebox.askyesno("تأكيد الحذف", f"هل أنت متأكد من حذف النشاط '{activity_name}'؟\nسيتم حذف سجلات الحضور المرتبطة به.", icon='warning', parent=self):
            self.db.delete_activity(activity_id)
            self.db.log_action(self.controller.current_user, "حذف نشاط", f"معرّف النشاط: {activity_id}")
            self.populate_table()

    def manage_attendance(self):
        selected_item = self.tree.focus()
        if not selected_item:
            messagebox.showwarning("تنبيه", "يرجى تحديد نشاط لإدارة الحضور.", parent=self)
            return
        
        values = self.tree.item(selected_item)['values']
        activity_id = values[-1]
        activity_name = values[-2]
        AttendanceDialog(self, self.controller, activity_id, activity_name)

# =================================================================
# نقطة انطلاق البرنامج
# =================================================================
if __name__ == "__main__":
    try:
        matplotlib.use("TkAgg")
    except ImportError:
        print("Matplotlib is not installed. Charts will not be available.")

    font_error = FontManager.check_and_register_fonts()
    if font_error:
        root = ctk.CTk()
        root.withdraw()
        messagebox.showerror("خطأ فادح - نقص في الخطوط", font_error)
        root.destroy()
        sys.exit(1)

    app = App()
    app.mainloop()
